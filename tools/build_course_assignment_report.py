from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "FocusFlow_Course_Assignment_Report.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEAL = RGBColor(15, 118, 110)
GRAY = RGBColor(86, 96, 110)
INK = RGBColor(24, 33, 47)
BORDER = "D9DEE7"
HEADER_FILL = "EEF3F8"
CALLOUT_FILL = "F4F6F9"


def set_run_font(run, size=11, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, align=None, before=0, after=8, line=1.333):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt({1: 18, 2: 12, 3: 8}.get(level, 8))
    p.paragraph_format.space_after = Pt({1: 10, 2: 6, 3: 4}.get(level, 4))
    run = p.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}.get(level, 11),
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}.get(level, BLUE),
        bold=True,
    )
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total = 0
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
            total += 0
            tc_pr = row.cells[i]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths)))

    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in (("top", 100), ("start", 140), ("bottom", 100), ("end", 140)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], HEADER_FILL)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=NAVY, size=9.2)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9.0 if len(value) < 44 else 8.6)
    add_para(doc, "", after=4)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, "D5DCE8")
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=TEAL, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=INK)
    add_para(doc, "", after=4)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for idx, size, color in [(1, 16, BLUE), (2, 13, BLUE), (3, 12, DARK_BLUE)]:
        style = doc.styles[f"Heading {idx}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.text = ""
    r = header.add_run("FocusFlow 人机交互课程作业报告")
    set_run_font(r, size=9, color=GRAY, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    r = footer.add_run("FocusFlow - Adaptive Attention Reading System")
    set_run_font(r, size=9, color=GRAY)


def cover(doc):
    add_para(doc, "Human-Computer Interaction Course Assignment", size=11, color=TEAL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=88, after=18)
    add_para(doc, "FocusFlow 课程作业报告", size=27, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "自适应注意力阅读管理系统的设计、原型与评估方案", size=15, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    add_para(doc, "本报告以课程作业为定位，重点说明问题分析、HCI 设计理由、原型方案、交互流程、评估设计与反思，而不是把文档写成开发维护手册。", size=11, color=GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
    add_table(doc, ["项目项", "课程报告中的表述重点"], [
        ["选题", "长文本阅读中的注意力漂移、理解停滞与恢复辅助。"],
        ["目标", "通过可运行原型说明一个“感知-判断-反馈-复盘”的人机交互闭环。"],
        ["方法", "结合用户场景分析、状态模型、低打扰交互设计和可评估实验方案。"],
        ["成果", "浏览器端原型、答辩 PPT、课程报告与后续实验设计。"],
    ], [1.35, 5.15])
    doc.add_page_break()


def overview(doc):
    add_heading(doc, "报告结构", 1)
    add_para(doc, "为避免内容被拆成过多零散小节，本版报告采用少数大章节组织。每个章节围绕一个课程报告问题展开：为什么选这个题，用户遇到什么问题，系统如何设计，交互如何成立，技术如何服务于体验，以及怎样评价这个 HCI 原型是否有效。")
    for item in [
        "第一章：选题背景与课程问题定位",
        "第二章：用户需求、使用场景与设计目标",
        "第三章：原型总体方案与 HCI 设计思路",
        "第四章：交互流程与界面体验设计",
        "第五章：关键实现与原型完成情况",
        "第六章：评估方案、数据指标与预期结果",
        "第七章：伦理、隐私、局限性与改进方向",
        "第八章：课程总结与个人/小组反思",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)
    doc.add_page_break()


CHAPTERS = [
    (
        "第一章 选题背景与课程问题定位",
        [
            ("选题来源", [
                "FocusFlow 的选题来自长文本阅读中的真实体验：很多学习任务并不是难在打开材料，而是难在持续保持有效阅读。学生阅读论文、报告、技术文档或课程材料时，经常会出现三种状态：已经离开阅读任务却没有意识到，停在同一段落迟迟无法理解，或者中断后不知道如何重新进入原来的阅读节奏。",
                "传统阅读器通常关注排版、笔记和进度记录，但它们很少尝试理解“用户此刻是否还在有效阅读”。因此，本课程作业把问题放在 HCI 的核心视角上：系统能不能通过可观察的行为信号理解用户状态，并以足够克制的方式帮助用户恢复阅读，而不是简单弹出粗暴提醒。"
            ]),
            ("课程作业定位", [
                "本项目不是要开发一个完整商业产品，也不是提交一份工程维护文档。它更接近一个 HCI 课程原型：用可运行系统承载一个交互设计假设，并通过系统结构、界面反馈和评估方案证明这个假设有被研究和验证的价值。",
                "因此，报告中的技术说明只服务于设计论证。摄像头、眼动、鼠标轨迹、滚动行为、状态机和摘要面板不是为了展示技术堆栈，而是为了回答课程问题：系统如何感知阅读状态，如何避免误判，如何选择干预强度，以及如何在不破坏沉浸感的前提下提供帮助。"
            ]),
        ],
        ("课程问题", "如何在浏览器端使用低成本、多模态交互信号推断用户的阅读状态，并通过低打扰、可解释、可恢复的反馈帮助用户回到阅读任务？"),
        ["把阅读辅助从静态内容呈现推进到动态状态理解。", "把“分心”与“理解困难”区分开，避免所有干预都变成提醒。", "强调用户控制权，使系统被理解为辅助工具而不是监控工具。"],
    ),
    (
        "第二章 用户需求、使用场景与设计目标",
        [
            ("用户与场景", [
                "目标用户主要是需要长时间阅读材料的学生和学习者，包括阅读课程论文、复习教材、浏览技术文档、准备答辩材料或进行在线学习的人。这个群体通常不缺少阅读工具，而是缺少一种能在阅读过程中温和提醒、辅助理解、并在结束后帮助复盘的系统。",
                "项目场景可以被概括为“长文本、较高认知负荷、需要保持连续注意力”的学习场景。用户可能在桌面浏览器中阅读，也可能在普通摄像头环境下尝试眼动追踪。考虑到课程展示和现实设备限制，系统必须允许摄像头不可用时继续以鼠标追踪模式运行。"
            ]),
            ("需求提炼", [
                "从用户角度看，系统需要做到三点。第一，它应该帮助用户意识到自己可能已经离开阅读任务；第二，它应该在用户理解困难时提供内容层面的辅助；第三，它应该在用户恢复阅读时减少压力，而不是不断提醒用户“你刚才分心了”。",
                "从课程评估角度看，系统需要提供可以被解释和测量的输出，例如状态时间线、注意力比例、阅读速度、分心次数、恢复时间和段落停留情况。这些指标使原型不仅能演示，还能进入后续用户实验。"
            ]),
        ],
        None,
        ["低打扰：优先使用弱提示，减少对阅读流的破坏。", "可解释：状态与干预原因应能被用户和答辩听众理解。", "可降级：摄像头不可用时仍可用鼠标与滚动信号完成演示。", "可评估：设计实验指标，为课程报告提供证据路径。"],
    ),
    (
        "第三章 原型总体方案与 HCI 设计思路",
        [
            ("总体闭环", [
                "FocusFlow 的总体方案不是一个单向监控系统，而是一个人机交互闭环。系统先感知用户行为，再推断当前阅读状态，然后选择合适的反馈策略，最后把反馈结果与会话数据记录下来，为用户复盘和后续评估提供材料。",
                "这个闭环的关键是“状态”这一中间层。系统不直接从某一个行为跳到某一个干预，而是先把行为信号解释为 Normal、Distracted、Struggling、Recovering 四类状态。这样做的好处是让反馈更有理由：分心对应的是回到任务，理解困难对应的是内容辅助，恢复阶段对应的是轻量鼓励。"
            ]),
            ("设计取舍", [
                "原型采用浏览器端实现，是为了降低课程展示成本。WebGazer 和摄像头用于眼动模式，鼠标与滚动用于降级模式，段落停留用于识别理解卡顿。多模态不是为了追求复杂，而是为了避免单一信号误判，例如用户短暂停止滚动并不一定代表分心，也可能是在认真理解。",
                "界面反馈遵循渐进干预原则。系统优先选择可忽略的视觉提示，只有在状态持续、置信度较高或连续多次出现时才提高干预强度。这样可以体现 HCI 中对用户自主性、注意力负担和交互节奏的关注。"
            ]),
        ],
        None,
        ["感知层：收集眼动、鼠标、滚动、段落停留等行为信号。", "认知层：将信号转换为四类可解释状态。", "决策层：根据状态、持续时间和严重程度选择干预。", "界面层：用视觉提示、摘要、进度反馈等方式呈现帮助。", "分析层：记录会话数据，支持复盘与实验评估。"],
    ),
    (
        "第四章 交互流程与界面体验设计",
        [
            ("用户流程", [
                "用户打开系统后，可以直接使用鼠标追踪模式阅读，也可以在允许摄像头权限后进入眼动追踪模式。用户导入文本或使用默认材料，开始阅读后系统持续观察交互信号。阅读过程中，用户不需要频繁操作系统，因为系统的主要职责是在后台判断状态，并在必要时以克制方式反馈。",
                "当系统判断用户可能分心时，界面可以给出弱视觉提示或浮动提示；当系统判断用户可能理解困难时，可以触发关键词高亮或段落概览；当用户恢复阅读时，系统可以显示进度或正向反馈。整个流程强调“帮助回到阅读”，而不是“打断并批评用户”。"
            ]),
            ("界面原则", [
                "界面设计需要避免把状态仪表盘放得过于强势。阅读内容应始终是视觉中心，状态、指标和日志只作为辅助信息存在。用户可以切换深度专注模式，以减少侧边栏和调试信息对阅读的干扰。",
                "段落级辅助是本项目与普通提醒工具的区别之一。系统不是只判断用户有没有分心，也会关注用户是否在某一段停留过久。对于学习场景来说，理解困难往往比单纯走神更重要，因此摘要、关键词和段落概览都被设计为理解支持，而不是注意力惩罚。"
            ]),
        ],
        None,
        ["阅读内容保持主位，分析信息保持次级。", "提示以可忽略、可恢复、可关闭为基本原则。", "摘要与关键词只辅助理解，不替代原文阅读。", "深度专注模式体现用户对系统可见程度的控制。"],
    ),
    (
        "第五章 关键实现与原型完成情况",
        [
            ("实现说明", [
                "本章只概括与课程作业相关的关键实现，不把文档写成开发手册。项目使用原生 JavaScript 组织模块，入口为 index.html，本地服务由 server/dev-server.js 提供。主要模块包括感知、认知、决策、界面、分析、NLP 和国际化。",
                "感知层整合人脸、视线区域、鼠标移动、滚动速度和段落停留。认知层用状态机维护四类状态，并记录状态历史。决策层用策略矩阵选择不同强度的干预。界面层负责阅读区域、视觉效果、专注模式和调试面板。分析层生成会话报告。"
            ]),
            ("完成情况", [
                "从课程原型角度看，FocusFlow 已经具备完整闭环：可以运行、可以阅读、可以采集交互信号、可以根据状态变化反馈，也可以在结束后生成报告。这说明原型不仅是静态界面展示，而是能够支撑课程答辩中的现场演示。",
                "需要强调的是，当前版本仍是研究原型。眼动精度、阈值设置、策略有效性和用户差异都需要后续实验验证。因此，本报告不会把系统描述成成熟产品，而是把它定位为一个可用于验证 HCI 假设的课程项目。"
            ]),
        ],
        None,
        ["可运行：本地服务启动后可在浏览器中使用。", "可降级：摄像头不可用时仍支持鼠标追踪。", "可解释：状态机输出可被展示和说明。", "可复盘：会话报告保留阅读过程中的关键指标。"],
    ),
    (
        "第六章 评估方案、数据指标与预期结果",
        [
            ("实验思路", [
                "为了证明系统有效，不能只依赖主观演示，而需要设计用户实验。建议比较三种条件：普通阅读器、鼠标追踪 FocusFlow、眼动追踪 FocusFlow。实验任务可以设置为阅读同等难度材料，并在阅读后完成理解题和主观问卷。",
                "实验关注的问题包括：FocusFlow 是否能缩短分心后的恢复时间，是否能提升阅读理解表现，是否会增加主观负担，用户是否接受系统干预，以及用户是否认为状态反馈具有解释性。"
            ]),
            ("评价指标", [
                "客观指标包括阅读理解正确率、任务完成时间、平均阅读速度、分心次数、恢复时间、状态转移次数和段落停留分布。主观指标包括 NASA-TLX 主观负担、系统可用性、感知打扰程度、干预接受度、隐私接受度和系统信任度。",
                "预期结果不是要求所有指标都显著提升。对于 HCI 课程作业来说，更重要的是形成合理的评估框架：如果系统提升理解但增加负担，需要讨论平衡；如果系统降低分心但用户觉得被监控，需要讨论伦理与控制权；如果鼠标模式效果接近眼动模式，则说明低成本信号有研究价值。"
            ]),
        ],
        None,
        ["普通阅读器用于建立基线。", "鼠标追踪模式验证低成本行为信号的价值。", "眼动追踪模式验证多模态感知是否带来额外收益。", "主观问卷用于检查系统是否真的被用户接受。"],
    ),
    (
        "第七章 伦理、隐私、局限性与改进方向",
        [
            ("伦理与隐私", [
                "注意力相关系统天然容易引发隐私担忧，因为它涉及摄像头、行为轨迹和用户状态判断。因此，FocusFlow 必须明确自己的定位：它是帮助用户恢复注意力的辅助系统，而不是判断用户是否认真的监控系统。",
                "在课程展示和后续实验中，应向用户说明摄像头用途、数据处理方式和日志保存范围。系统应提供关闭追踪、清除日志、使用鼠标模式和退出干预的入口。只有当用户知道系统在做什么、为什么做、如何关闭时，状态感知才具有合理性。"
            ]),
            ("局限与改进", [
                "当前原型的主要局限包括浏览器端眼动精度不稳定、不同用户阅读节奏差异较大、状态阈值尚未经过实验校准、摘要辅助可能存在延迟或不准确，以及强提示可能破坏阅读沉浸感。",
                "后续改进可以集中在四个方向：个性化阈值学习、段落级理解辅助质量提升、实验数据导出与统计分析、以及更透明的隐私设置界面。对于课程项目而言，这些方向不一定都要实现，但需要在报告中体现研究延展性。"
            ]),
        ],
        None,
        ["不保存原始摄像头视频，优先使用即时计算结果。", "允许用户选择鼠标模式，降低摄像头依赖。", "所有干预都应可关闭、可忽略、可恢复。", "报告中应承认误判风险，而不是把状态机描述成绝对准确。"],
    ),
    (
        "第八章 课程总结与反思",
        [
            ("课程收获", [
                "通过 FocusFlow，本课程作业把 HCI 的多个核心概念连接起来：用户需求、认知负荷、注意力、可解释反馈、低打扰交互、隐私边界和可评估原型。项目的价值不在于实现了多少功能，而在于用一个可运行系统表达了一个清晰的人机交互问题。",
                "这个项目也说明，原型不需要一开始就成为完整产品。课程作业中的原型更重要的是能够让设计假设被看见、被操作、被讨论和被评估。FocusFlow 的感知层、状态机和干预策略共同构成了这个假设的可视化表达。"
            ]),
            ("反思", [
                "如果继续完善，本项目最需要补充的是真实用户实验。当前系统已经能够展示交互闭环，但还不能证明所有干预都有效，也不能证明用户一定接受这种状态感知。因此，下一步应从小规模课堂实验开始，收集主观反馈和行为数据，再调整阈值与反馈方式。",
                "从课程报告角度看，本版文档把重点放在“为什么这样设计”和“如何评价这个设计”上，而不是罗列过多开发细节。这样更符合人机交互课程作业的表达方式，也更便于老师理解项目的研究问题和设计价值。"
            ]),
        ],
        ("最终结论", "FocusFlow 是一个围绕长文本阅读注意力管理展开的 HCI 课程原型。它通过多模态行为感知、可解释状态模型和低打扰反馈，展示了一个从用户问题到交互方案再到评估设计的完整课程作业闭环。"),
        ["问题明确：针对长文本阅读中的分心、卡顿和恢复困难。", "方案完整：包含感知、状态、干预、报告四个关键环节。", "表达适合课程作业：重点在 HCI 设计论证，而不是工程开发堆栈。"],
    ),
]


def add_chapter(doc, title, sections, callout, bullets):
    add_heading(doc, title, 1)
    for sub_title, paragraphs in sections:
        add_heading(doc, sub_title, 2)
        for text in paragraphs:
            add_para(doc, text)
    if callout:
        add_callout(doc, callout[0], callout[1])
    if bullets:
        add_heading(doc, "本章小结", 2)
        for item in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(item)
            set_run_font(r, size=10.5, color=INK)
    if title.startswith("第二章"):
        add_table(doc, ["需求类别", "课程报告中的含义", "对应设计"], [
            ["注意力恢复", "用户离开阅读任务后需要温和回到原文。", "弱提示、浮动提示、恢复反馈。"],
            ["理解支持", "用户长时间停留可能代表内容难以理解。", "关键词高亮、段落概览、摘要面板。"],
            ["用户控制", "用户应能决定是否启用摄像头和是否查看分析信息。", "鼠标模式、深度专注模式、可关闭提示。"],
            ["实验验证", "课程作业需要说明如何评价系统是否有效。", "状态时间线、理解题、主观问卷。"],
        ], [1.25, 2.55, 2.7])
    if title.startswith("第三章"):
        add_table(doc, ["环节", "设计含义", "课程展示重点"], [
            ["感知", "从行为中获得阅读状态线索。", "说明使用多模态信号而非单一判断。"],
            ["判断", "把信号组织成可解释状态。", "展示 Normal、Distracted、Struggling、Recovering。"],
            ["反馈", "根据状态选择低打扰干预。", "展示策略矩阵与渐进式反馈。"],
            ["复盘", "把阅读过程变成可讨论的数据。", "展示会话报告和状态时间线。"],
        ], [1.1, 2.6, 2.8])
    if title.startswith("第五章"):
        add_table(doc, ["模块", "在课程作业中的作用", "不展开为开发细节的原因"], [
            ["Perception", "说明系统如何感知用户行为。", "重点是信号来源和设计意义，不是逐行代码。"],
            ["StateMachine", "说明如何把行为转换为认知状态。", "重点是可解释状态模型。"],
            ["Intervention", "说明干预为何分级。", "重点是低打扰原则。"],
            ["SessionReport", "说明如何支持复盘与评估。", "重点是课程实验指标。"],
        ], [1.45, 2.6, 2.45])
    if title.startswith("第六章"):
        add_table(doc, ["指标", "测量方式", "解释意义"], [
            ["阅读理解正确率", "阅读后答题。", "判断辅助是否损害或提升理解。"],
            ["恢复时间", "从 Distracted 到 Recovering/Normal 的时间。", "判断系统是否帮助用户回到任务。"],
            ["主观负担", "NASA-TLX 或简化问卷。", "判断干预是否增加压力。"],
            ["干预接受度", "Likert 量表与访谈。", "判断用户是否愿意长期使用。"],
        ], [1.55, 2.35, 2.6])
    doc.add_page_break()


def main():
    doc = Document()
    configure(doc)
    cover(doc)
    overview(doc)
    for chapter in CHAPTERS:
        add_chapter(doc, *chapter)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
