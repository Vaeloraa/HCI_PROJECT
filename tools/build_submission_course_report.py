# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "FocusFlow_Submission_Course_Report.docx"

INK = RGBColor(28, 36, 48)
MUTED = RGBColor(90, 101, 116)
NAVY = RGBColor(20, 52, 88)
BLUE = RGBColor(43, 95, 151)
TEAL = RGBColor(20, 125, 122)
RED = RGBColor(165, 65, 54)
LINE = "D7DEE8"
HEADER_FILL = "EDF3F8"
SOFT_FILL = "F7F9FC"
TITLE_FILL = "E9F0F7"


def set_run_font(run, size=10.5, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_format(paragraph, before=0, after=6, line=1.5, first_line=True):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)


def add_para(doc: Document, text: str = "", size=12, color=INK, bold=False, italic=False,
             align=None, before=0, after=6, line=1.5, first_line=True):
    paragraph = doc.add_paragraph()
    set_para_format(paragraph, before=before, after=after, line=line, first_line=first_line)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_rich_para(doc: Document, parts: Sequence[tuple[str, bool]], size=12, before=0,
                  after=6, line=1.5, first_line=True):
    paragraph = doc.add_paragraph()
    set_para_format(paragraph, before=before, after=after, line=line, first_line=first_line)
    for text, bold in parts:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=INK, bold=bold)
    return paragraph


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt({1: 17, 2: 11, 3: 7}.get(level, 6))
    paragraph.paragraph_format.space_after = Pt({1: 8, 2: 5, 3: 3}.get(level, 3))
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size={1: 15.5, 2: 12.5, 3: 11.5}.get(level, 10.5),
        color={1: NAVY, 2: BLUE, 3: TEAL}.get(level, BLUE),
        bold=True,
    )
    return paragraph


def set_cell_text(cell, text: str, bold=False, color=INK, size=9.2):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=LINE):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[float]):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]],
              widths: Sequence[float], font_size=8.9):
    rows = list(rows)
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, HEADER_FILL)
        set_cell_text(cell, header, bold=True, color=NAVY, size=font_size)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            if r_idx % 2 == 0:
                shade_cell(cell, SOFT_FILL)
            set_cell_text(cell, text, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_quote(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table, color="C8D2E0")
    set_table_geometry(table, [6.75])
    cell = table.rows[0].cells[0]
    shade_cell(cell, TITLE_FILL)
    set_cell_text(cell, text, bold=False, color=NAVY, size=9.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.first_line_indent = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("FocusFlow 人机交互课程作业  |  第 ")
    set_run_font(run, size=8.5, color=MUTED)
    page_run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    page_run._r.append(fld_char1)
    page_run._r.append(instr_text)
    page_run._r.append(fld_char2)
    run = footer.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.start_type = WD_SECTION.CONTINUOUS
    add_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for level in (1, 2, 3):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.bold = True
        style.font.color.rgb = {1: NAVY, 2: BLUE, 3: TEAL}[level]


def add_cover(doc: Document):
    add_para(doc, "FocusFlow 自适应注意力阅读管理系统课程作业",
             size=21, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=8, after=8, line=1.05, first_line=False)
    add_para(doc, "基于 WebGazer.js、行为感知与自适应干预的人机交互系统分析",
             size=12, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=0, after=10, first_line=False)

    add_table(
        doc,
        ["项目项", "内容"],
        [
            ["课程方向", "Human-Computer Interaction / 人机交互课程作业"],
            ["项目名称", "FocusFlow"],
            ["项目仓库", "Vaeloraa/HCI_PROJECT"],
            ["核心技术", "WebGazer.js、浏览器行为感知、有限状态机、自适应阈值、阅读辅助与会话分析"],
            ["提交日期", "2026 年 6 月 16 日"],
        ],
        [1.55, 5.2],
        font_size=9.4,
    )

    add_para(doc, "摘要", size=12.5, color=NAVY, bold=True, before=8, after=4, first_line=False)
    add_para(
        doc,
        "FocusFlow 面向长文本在线阅读中的注意力漂移、理解停滞和回到任务困难等典型人机交互问题，"
        "通过浏览器端感知模块采集鼠标、滚动、凝视、头部姿态与页面位置等信号，再由认知状态机将连续行为解释为 Normal、"
        "Distracted、Struggling、Recovering 四类状态，最后由决策模块选择低打扰的界面干预。项目不是单纯的阅读器，也不是独立的眼动演示，"
        "而是一个把感知、认知判断、干预反馈和结果分析连接在同一闭环中的自适应注意力管理系统。",
        after=4,
    )
    add_para(
        doc,
        "代码层面，系统以 index.html 提供主界面，以 js/main.js 作为应用控制器，以 js/perception、js/cognition、"
        "js/decision、js/ui、js/nlp、js/analytics 等目录形成分层结构。实际运行默认使用鼠标与滚动行为作为稳定入口，"
        "在用户授权摄像头后切换到 gaze 模式并启用 WebGazer.js 与九点校准流程。该设计降低了硬件依赖和权限失败对课程演示的影响，"
        "也保留了眼动追踪在人机交互研究中的扩展价值。",
        after=4,
    )
    add_para(
        doc,
        "关键词：FocusFlow；人机交互；注意力管理；WebGazer.js；有限状态机；自适应界面；阅读辅助",
        size=9.8,
        color=MUTED,
        bold=True,
        first_line=False,
    )


def add_section_1(doc: Document):
    add_heading(doc, "1 项目背景与问题定义", 1)
    add_heading(doc, "1.1 长文本阅读中的注意力问题", 2)
    add_heading(doc, "1.1.1 注意力漂移不是单一瞬间事件", 3)
    add_para(
        doc,
        "在线阅读的注意力下降通常不是一个突然发生的动作，而是由多种细小行为逐步累积形成的状态。"
        "用户可能先停止滚动，随后鼠标长时间不移动，接着视线或指针在页面边缘停留，最后才表现为真正离开阅读任务。"
        "如果系统只在用户完全离开页面后提醒，干预已经滞后；如果系统在任何短暂停顿时都提醒，又会破坏正常阅读节奏。"
        "FocusFlow 将注意力变化看作连续状态，而不是简单的二元在线/离线判断。"
    )
    add_para(
        doc,
        "项目中这种判断体现在 js/perception/perceptionModule.js 对 interactionActive、dwellTime、"
        "scrollVelocity、scrollIdleDuration、mouseIdleDuration 等特征的组合计算上。"
        "这些特征不是为了单独得出结论，而是交给 js/cognition/stateMachine.js 进行概率化融合。"
        "因此，一次停顿不会直接被判为走神，只有停顿时长、交互缺失、凝视区域分散等证据共同增强时，Distracted 概率才会上升。"
    )

    add_heading(doc, "1.1.2 理解停滞需要与正常精读区分", 3)
    add_para(
        doc,
        "用户在难句、数据表或概念密集段落前停留较久，并不一定代表失败；很多时候这是主动精读。"
        "人机交互系统如果把所有长停留都看作问题，会把高质量阅读误判成障碍。"
        "FocusFlow 的处理方式是把 dwellTime 与滚动速度、滚动暂停、鼠标阅读轨迹、方向变化等特征一起使用："
        "当停留超过 strugglingDwellTime，且滚动速度低、滚动长时间暂停或鼠标移动缺少连续阅读模式时，"
        "Struggling 才会被认为更可能发生。"
    )
    add_para(
        doc,
        "这种设计符合课程中以人为中心的原则：系统不是用技术替代用户判断，而是在用户可能需要帮助时提供候选帮助。"
        "帮助形式也保持可撤销和低压力，例如关键词高亮、段落摘要按钮、理解卡片等，而不是强制跳转页面或打断阅读。"
    )

    add_heading(doc, "1.1.3 恢复过程也是交互设计对象", 3)
    add_para(
        doc,
        "许多注意力工具只关注检测分心，却忽略用户回到任务后的恢复阶段。"
        "FocusFlow 在状态机中单独设置 Recovering 状态，用来描述用户从 Distracted 或 Struggling 回到稳定阅读前的过渡过程。"
        "这种状态的存在很重要，因为恢复期不适合继续升级打扰，反而适合用轻量进度、积极反馈或界面降噪来帮助用户重新建立节奏。"
    )
    add_para(
        doc,
        "代码中 Recovering 的证据包括面部重新出现、交互恢复、滚动恢复以及当前已处于 Recovering 的持续证据。"
        "这些条件让系统能识别“用户正在回来”，从而避免把刚恢复的用户再次推入强提醒。"
        "从交互体验看，这比单纯的提醒系统更接近自适应界面，因为它对用户状态的变化方向也作出响应。"
    )

    add_heading(doc, "1.2 课程作业目标与设计边界", 2)
    add_heading(doc, "1.2.1 作业目标", 3)
    add_para(
        doc,
        "FocusFlow 的目标是在普通浏览器环境中实现一个可演示、可解释、可扩展的自适应阅读系统。"
        "可演示意味着系统不依赖昂贵眼动仪，使用 WebGazer.js 与鼠标滚动特征即可完成课堂展示；"
        "可解释意味着每一种状态判断和干预策略都可以从代码中的阈值、权重与事件流追溯；"
        "可扩展意味着后续可以把更准确的模型、更多材料类型或更细粒度的学习分析加入现有框架。"
    )
    add_para(
        doc,
        "项目没有把目标设定为医学诊断、心理评估或强制监控，而是限定在学习阅读场景的即时辅助。"
        "这一区分对课程作业很关键：系统可以分析注意力相关行为，但不应将用户标签化；可以保存会话指标，"
        "但应避免收集超出阅读体验所需的数据；可以使用摄像头，但必须通过权限提示、默认鼠标模式和可退出机制降低隐私压力。"
    )

    add_heading(doc, "1.2.2 人机交互价值", 3)
    add_para(
        doc,
        "从 HCI 角度看，FocusFlow 的价值不在于某一个算法单点，而在于把用户状态建模和界面反馈放在同一条交互链路中。"
        "感知层观察行为，认知层解释状态，决策层选择干预，界面层呈现反馈，分析层记录效果。"
        "这条链路让系统能够回答三个课程层面的问题：如何感知用户状态，如何把状态变化转换成可解释的界面行为，"
        "以及如何在提供帮助时维持用户的控制感。"
    )
    add_quote(
        doc,
        "核心交互命题：系统应当在用户注意力变弱时轻轻托住阅读流程，而不是抢走阅读任务本身。"
    )


def add_section_2(doc: Document):
    add_heading(doc, "2 需求分析与设计目标", 1)
    add_heading(doc, "2.1 用户场景", 2)
    add_heading(doc, "2.1.1 学习者阅读课程材料", 3)
    add_para(
        doc,
        "主要用户是需要阅读较长课程材料、论文节选、实验说明或项目文档的学生。"
        "这类用户的阅读任务具有持续时间长、信息密度高、反馈延迟明显的特点。"
        "普通阅读器只能显示文本，无法判断用户是否停在难点、是否离开任务、是否需要局部解释。"
        "FocusFlow 将文本按段落块映射到界面元素中，使系统能够把用户行为与具体文本块关联起来。"
    )
    add_para(
        doc,
        "在实际代码里，js/ui/readingView.js 负责构建阅读块、计算词数、维护当前块和 WPM 等指标；"
        "js/ui/paragraphSplitter.js 负责把导入文本拆成标题、段落和长段落切片。"
        "这种结构保证交互分析不是停留在整个页面层面，而是能落到“用户正在阅读第几个 block、该 block 的文本是什么、是否需要摘要或高亮”的粒度上。"
    )

    add_heading(doc, "2.1.2 摄像头不可用时的课程演示场景", 3)
    add_para(
        doc,
        "课堂演示常见问题包括浏览器权限被拒、摄像头被占用、光线不足、WebGazer 模型初始不稳定等。"
        "FocusFlow 默认进入 mouse 模式，并用鼠标、滚动和停留特征驱动状态机。"
        "当用户主动启用 gaze 模式时，系统才请求摄像头并进入 WebGazer 初始化和校准流程。"
        "这种默认策略提高了演示可靠性，也让项目更符合隐私最小化原则。"
    )
    add_para(
        doc,
        "js/utils/cameraAccess.js 对摄像头支持、安全上下文、权限请求、释放和错误分类进行了封装。"
        "js/main.js 中的 trackingMode 默认值是 'mouse'，并通过 focusflow-tracking-mode 事件向界面同步状态。"
        "如果 gaze 初始化失败，应用可以回到 mouse 模式继续运行，而不是让整个课程项目无法展示。"
    )

    add_heading(doc, "2.2 功能需求", 2)
    add_table(
        doc,
        ["需求类别", "具体需求", "代码实现位置", "交互意义"],
        [
            ["阅读材料呈现", "加载默认文本或用户导入文本，拆分为可定位的阅读块。", "js/ui/readingView.js；js/ui/paragraphSplitter.js", "让行为数据能够映射到具体段落，支撑段落级辅助。"],
            ["行为感知", "采集鼠标移动、点击、滚动、凝视坐标、面部与头部姿态等信号。", "js/perception/perceptionModule.js 及其子模块", "形成状态判断所需的多源证据，降低单一传感器误判。"],
            ["状态判断", "将连续行为解释为 Normal、Distracted、Struggling、Recovering。", "js/cognition/stateMachine.js", "把低层数据转换成可被界面使用的用户状态。"],
            ["自适应干预", "根据状态、置信度、持续时间和冷却时间选择提示、高亮、摘要或反馈。", "js/decision/decisionModule.js；js/decision/interventionStrategy.js", "使帮助与用户状态匹配，减少无意义提醒。"],
            ["理解辅助", "对稳定阅读块预取摘要，必要时显示理解卡片或关键词。", "js/nlp/llmSummaryManager.js；js/nlp/paragraphSummarizer.js；js/nlp/keywordExtractor.js", "把困难段落转化为更易处理的信息单元。"],
            ["结果分析", "记录状态变化、阅读速度、专注比例、分心恢复、热力图和会话报告。", "js/analytics/attentionAnalytics.js；js/analytics/sessionReport.js", "为评估系统效果和课堂答辩提供证据。"],
        ],
        [1.15, 2.0, 1.7, 1.9],
    )

    add_heading(doc, "2.3 非功能需求与伦理要求", 2)
    add_heading(doc, "2.3.1 可解释性", 3)
    add_para(
        doc,
        "课程作业需要能够解释系统为什么在某个时刻提醒用户。"
        "FocusFlow 没有把关键判断完全交给黑箱模型，而是在状态机中保留了清晰的证据权重、阈值和转换条件。"
        "例如 Distracted 的判断会查看 gaze 模式下的 faceAbsentDuration、交互缺失、idleDuration、mouse 模式下的长时间不动以及热力图 dispersion。"
        "这些字段可以在调试面板和代码审阅中被追踪，使系统行为具备可辩护性。"
    )
    add_heading(doc, "2.3.2 低打扰性", 3)
    add_para(
        doc,
        "注意力辅助工具最容易出现的问题是过度提醒。"
        "项目通过策略冷却时间、全局最小干预间隔、同一状态连续次数驱动的升级等级、以及恢复状态的轻量反馈来控制打扰强度。"
        "js/decision/interventionStrategy.js 中至少保留 4000ms 的全局策略间隔，并为每种策略设置单独 cooldown，"
        "从代码层面限制界面频繁弹出。"
    )
    add_heading(doc, "2.3.3 隐私与权限", 3)
    add_para(
        doc,
        "摄像头只在用户选择 gaze 模式时使用，系统默认不依赖摄像头即可运行。"
        "这符合数据最小化原则，也让用户能够在不暴露面部信息的情况下体验主要功能。"
        "server/dev-server.js 中的 LLM 摘要代理只处理摘要请求，不承担用户身份认证或长期画像存储；"
        "前端在 LLM 不可用时使用规则摘要回退，使核心阅读辅助不被外部服务完全绑定。"
    )


def add_section_3(doc: Document):
    add_heading(doc, "3 项目框架与代码结构分析", 1)
    add_heading(doc, "3.1 运行入口与资源组织", 2)
    add_heading(doc, "3.1.1 package.json 与运行方式", 3)
    add_para(
        doc,
        "项目根目录的 package.json 将应用定义为 focusflow，描述为基于 WebGazer.js 的自适应注意力管理系统。"
        "脚本 serve 与 start 都指向 node server/dev-server.js，说明项目的推荐运行方式是通过本地 Node 服务提供静态文件与摘要接口。"
        "依赖项包含 webgazer ^3.5.3 与 chart.js ^4.4.0，分别对应眼动追踪和图表可视化能力。"
    )
    add_para(
        doc,
        "server/dev-server.js 兼具静态资源服务和 LLM 摘要代理功能。"
        "它读取 .env 中的 OPENAI_API_KEY、OPENAI_BASE_URL、OPENAI_MODEL 等配置，提供 /api/llm/status 与 /api/summarize 两个接口。"
        "当 API key 不存在时，服务返回明确状态，前端可以进入 fallback 模式。"
        "这让项目在无密钥环境下仍可展示基础功能，在有密钥环境下增强段落理解辅助。"
    )

    add_heading(doc, "3.1.2 index.html 与前端装配", 3)
    add_para(
        doc,
        "index.html 不是简单的静态页面，而是承担了阅读器界面、控制按钮、状态展示、主题切换、凝视光标、计时器和多语言文本挂载等职责。"
        "页面中存在 ff-reading-content、ff-camera-gate、ff-tracking-toggle、ff-timer-display、ff-gaze-cursor 等关键 DOM 节点，"
        "这些节点被 js/main.js 和 UI 模块调用。"
    )
    add_para(
        doc,
        "页面脚本通过 DOMContentLoaded、focusflow-state-change、focusflow-tracking-mode、focusflow-lang-change、beforeunload 等事件组织生命周期。"
        "这种事件式连接方式让状态机和界面渲染之间保持松耦合：状态机不直接操作所有视觉元素，而是发出状态变化；"
        "界面层再根据事件更新文字、图标、提示和效果。"
    )

    add_heading(doc, "3.2 主控制器 main.js 的闭环流程", 2)
    add_para(
        doc,
        "js/main.js 是项目最核心的编排文件，当前文件长度超过 1400 行，承担模块初始化、跟踪模式切换、校准流程、感知更新、状态更新、策略执行、"
        "摘要预取、会话报告展示和资源销毁等职责。它在初始化阶段依次创建 ReadingView、PerceptionModule、StateMachine、DecisionModule、"
        "VisualEffects、FocusMode、DebugPanel、KeywordExtractor、AttentionAnalytics 和 LLMSummaryManager。"
    )
    add_para(
        doc,
        "主流程可以概括为：用户阅读时产生鼠标、滚动或凝视数据；PerceptionModule 将这些数据转换为特征向量；"
        "StateMachine 根据特征更新状态概率；DecisionModule 结合状态、用户画像和阈值给出策略；"
        "VisualEffects 与 ReadingView 把策略转化为界面变化；AttentionAnalytics 记录行为与结果。"
        "这条路径在 main.js 的 gaze 更新和 runCognitionTick 中都能看到，前者处理实时凝视数据，后者保证没有新数据时状态持续时间仍然更新。"
    )
    add_quote(
        doc,
        "核心数据流：Calibration / WebGazer / Mouse → Perception → Cognition → Decision → UI Effects → Analytics。"
    )

    add_heading(doc, "3.3 分层架构", 2)
    add_table(
        doc,
        ["层次", "主要文件", "责任", "与其他层的关系"],
        [
            ["入口与服务层", "package.json；server/dev-server.js；index.html", "启动应用、提供静态资源、代理摘要接口、承载页面结构。", "向前端模块提供运行环境和必要接口。"],
            ["主控制层", "js/main.js", "统一初始化模块，切换 mouse/gaze 模式，连接状态事件、策略执行和会话报告。", "把感知、认知、决策、界面、分析层串成闭环。"],
            ["感知层", "js/perception/*.js；js/calibration/calibrationManager.js", "处理 WebGazer、面部、头部、鼠标、滚动、热力图和九点校准。", "输出 features 给认知层。"],
            ["认知层", "js/cognition/stateMachine.js", "用证据权重、状态概率、冷却和滞后来识别阅读状态。", "接收 features，输出当前状态与置信度。"],
            ["决策层", "js/decision/*.js", "根据状态选择干预策略，并持续适配阈值。", "接收 state/features，输出策略对象。"],
            ["界面反馈层", "js/ui/*.js", "阅读布局、段落切分、视觉效果、专注模式和调试面板。", "把策略转为可见且可控的交互反馈。"],
            ["NLP 辅助层", "js/nlp/*.js", "关键词提取、规则摘要、LLM 摘要管理与缓存。", "为 Struggling 状态和段落辅助提供内容支持。"],
            ["分析层", "js/analytics/*.js", "记录状态、速度、热力图、理解辅助使用和会话报告。", "为系统评估和迭代提供数据。"],
            ["工具层", "js/utils/*.js", "摄像头权限、文本编码识别等通用能力。", "减少入口层和感知层的异常处理复杂度。"],
        ],
        [1.0, 1.55, 2.2, 1.95],
    )

    add_heading(doc, "3.4 代码规模与课程完成度", 2)
    add_para(
        doc,
        "从代码规模看，项目已经超出简单静态页面：webgazer.js 作为第三方库占据最大体积，业务代码则分布在 main.js、ReadingView、"
        "VisualEffects、AttentionAnalytics、CalibrationManager、StateMachine、DecisionModule 等文件中。"
        "这种规模说明项目完成的不只是界面展示，还包含行为处理、状态建模、策略决策、摘要服务和分析输出。"
    )
    add_para(
        doc,
        "从课程作业角度看，这种分层结构的优点是答辩时可以清楚说明每个成员或模块的贡献：感知与认知模块回答“如何知道用户状态”，"
        "自适应界面模块回答“如何把状态变成低打扰反馈”，NLP 与分析模块回答“如何帮助理解并评估效果”。"
        "项目框架具有完整性，但 main.js 和 VisualEffects 仍然承担了较多职责，后续若继续迭代，可以把策略执行、摘要按钮状态和可视化更新进一步拆分。"
    )


def add_section_4(doc: Document):
    add_heading(doc, "4 感知层实现分析", 1)
    add_heading(doc, "4.1 PerceptionModule 的聚合职责", 2)
    add_heading(doc, "4.1.1 多源行为数据接入", 3)
    add_para(
        doc,
        "js/perception/perceptionModule.js 是感知层的统一入口。它创建 FaceDetection、HeadPoseEstimation、GazeRegionMapping、"
        "MouseTracker、ScrollAnalyzer 和 AttentionHeatmap，并监听 mousemove、click、scroll、wheel 等浏览器事件。"
        "这种设计把不同传感器和交互事件收束到一个模块中，避免状态机直接依赖 DOM 事件细节。"
    )
    add_para(
        doc,
        "getFeatures() 是感知层对外最重要的接口。它会根据当前 trackingMode 判断是否启用 faceTracking，"
        "在 gaze 模式下使用 facePresent 和 gaze 区域信息，在 mouse 模式下则把鼠标空闲时间作为保守替代。"
        "这让同一个认知状态机可以同时服务鼠标模式和眼动模式，而不需要为两种输入路径写两套状态逻辑。"
    )

    add_heading(doc, "4.1.2 特征向量的设计", 3)
    add_para(
        doc,
        "特征向量既包含瞬时行为，也包含持续时间。瞬时行为如 scrollVelocity、mouseVelocity、gazeRegion；"
        "持续时间如 idleDuration、faceAbsentDuration、dwellTime、scrollIdleDuration。"
        "对注意力系统来说，持续时间比单次坐标更有解释价值，因为用户的阅读状态通常体现在一段时间的稳定模式中。"
    )
    add_para(
        doc,
        "PerceptionModule 还会派生 interactionActive，用于表示用户是否仍在进行有效交互。"
        "该字段不是简单等于鼠标是否移动，而是综合 liveIsMoving、liveIsScrolling、liveMouseIdle < 3000 等条件。"
        "这种派生字段简化了认知层的判断，使状态机可以直接把 interactionActive 作为 Normal、Distracted 或 Recovering 的证据。"
    )

    add_heading(doc, "4.2 鼠标与滚动行为分析", 2)
    add_heading(doc, "4.2.1 MouseTracker", 3)
    add_para(
        doc,
        "js/perception/mouseTracker.js 记录鼠标位置、速度、方向变化和移动模式，并尝试区分 reading 与 searching 等行为。"
        "阅读行为往往具有横向移动比例较高、速度较平缓、方向变化不过度杂乱的特征；搜索或游移行为则可能出现更大幅度和更高频方向变化。"
        "状态机在 Normal 判断中使用 mouse reading horizontal movement > 0.6 作为弱证据，在 Struggling 判断中使用 directionChanges > 5 且缺少阅读模式作为困难证据。"
    )
    add_para(
        doc,
        "鼠标特征的意义在于它不需要额外权限，适合课程项目的默认模式。"
        "虽然鼠标并不等同于眼睛，但在网页阅读中，用户常把鼠标作为临时指针、定位工具或阅读辅助。"
        "因此，鼠标模式能提供一种低成本的注意力近似判断。"
    )

    add_heading(doc, "4.2.2 ScrollAnalyzer", 3)
    add_para(
        doc,
        "js/perception/scrollAnalyzer.js 分析滚动速度、滚动方向、滚动空闲、暂停、爆发式滚动和阅读节奏。"
        "滚动速度非常低并不总是坏事，但当低滚动速度与长 dwellTime、长 scrollIdleDuration 同时出现时，"
        "它会增强 Struggling 的可能性。相反，恢复滚动可以成为 Recovering 的证据。"
    )
    add_para(
        doc,
        "滚动信号在长文本阅读中特别重要，因为它直接体现用户在文档中的推进节奏。"
        "FocusFlow 没有用固定阅读速度要求用户，而是把滚动作为状态证据之一，再由 AdaptiveThreshold 根据个人阅读速度和停留习惯调整阈值。"
    )

    add_heading(doc, "4.3 眼动、面部、头部与校准", 2)
    add_heading(doc, "4.3.1 WebGazer 与九点校准", 3)
    add_para(
        doc,
        "gaze 模式由 js/main.js 触发 WebGazer 初始化，并通过 js/calibration/calibrationManager.js 提供九点校准流程。"
        "CalibrationManager 在全屏覆盖层中展示 3×3 网格点，采集每个点的 gaze sample，并在完成后把数据保存到 window.__focusflow_calibration。"
        "九点校准使浏览器端眼动估计更适合当前用户和屏幕环境，降低 gaze cursor 与真实阅读位置之间的偏差。"
    )
    add_para(
        doc,
        "项目还引入 js/perception/kalmanFilter.js 中的 KalmanFilter2D 对凝视坐标进行平滑。"
        "浏览器端眼动估计容易受头部微动、光照和模型误差影响，直接使用原始坐标会导致高亮跳动。"
        "卡尔曼滤波通过预测与观测融合减小噪声，使阅读块定位、凝视光标和热力图更新更稳定。"
    )

    add_heading(doc, "4.3.2 FaceDetection、HeadPose 与 GazeRegion", 3)
    add_para(
        doc,
        "js/perception/faceDetection.js 维护 facePresent、confidence、bounding box 和 faceAbsentDuration。"
        "在 gaze 模式下，faceAbsentDuration 是 Distracted 判断的重要证据；在 mouse 模式下，系统不强行要求面部存在，避免无摄像头时误判。"
    )
    add_para(
        doc,
        "js/perception/headPose.js 估计头部方向，并提供 getGazeRegion 等辅助判断。"
        "js/perception/gazeRegion.js 则把坐标映射到页面区域，维护 dwell、average dwell、most viewed region 和 trail。"
        "这些模块共同把原始点坐标转化为可以解释的阅读位置和停留模式。"
    )

    add_heading(doc, "4.4 特征字段与状态含义", 2)
    add_table(
        doc,
        ["特征", "来源模块", "主要含义", "参与判断"],
        [
            ["facePresent / faceAbsentDuration", "FaceDetection；PerceptionModule", "用户面部是否可见以及不可见持续时间。", "gaze 模式下增强 Distracted；面部返回增强 Recovering。"],
            ["interactionActive / idleDuration", "PerceptionModule", "鼠标和滚动是否仍在进行有效交互。", "活跃增强 Normal 或 Recovering；长时间不活跃增强 Distracted。"],
            ["dwellTime", "GazeRegionMapping；ScrollAnalyzer", "用户在当前区域或段落附近停留的时间。", "适度停留可支持 Normal；超长停留增强 Struggling。"],
            ["scrollVelocity / scrollIdleDuration", "ScrollAnalyzer", "文本推进速度与滚动暂停时间。", "平稳滚动支持 Normal；长期暂停结合 dwell 支持 Struggling。"],
            ["mouseVelocity / horizontalRatio", "MouseTracker", "鼠标速度与横向阅读式移动比例。", "横向阅读轨迹支持 Normal；杂乱方向变化可能支持 Struggling。"],
            ["dispersion / heatmap", "AttentionHeatmap", "注意力点在页面上的分散程度与热区。", "过度分散可作为 Distracted 的弱证据。"],
            ["gazeBlock index", "ReadingView；main.js", "凝视或指针所在阅读块。", "用于摘要预取、段落高亮、WPM 和当前块记录。"],
        ],
        [1.35, 1.55, 2.1, 1.75],
    )
    add_para(
        doc,
        "感知层的关键优点是没有把某一个字段绝对化。例如 faceAbsentDuration 只在 gaze 模式中具有强意义，"
        "而 mouse 模式更依赖 idleDuration 和 scrollIdleDuration。"
        "这种模式区分使项目可以同时兼顾演示稳定性和眼动研究价值。"
    )


def add_section_5(doc: Document):
    add_heading(doc, "5 认知状态建模分析", 1)
    add_heading(doc, "5.1 状态定义", 2)
    add_para(
        doc,
        "js/cognition/stateMachine.js 定义了 Normal、Distracted、Struggling、Recovering 四种核心状态。"
        "Normal 代表阅读节奏基本稳定；Distracted 代表用户可能离开任务或明显缺少交互；"
        "Struggling 代表用户仍在阅读上下文中，但可能卡在某个段落或概念上；Recovering 代表用户正在从分心或困难中恢复。"
    )
    add_para(
        doc,
        "四状态模型比二分类更适合阅读辅助。若只有专注和分心两类，系统无法区分“用户走神”和“用户读不懂”。"
        "这两种状态需要不同干预：走神更适合轻提醒或界面唤回，读不懂更适合摘要、关键词或解释。"
        "Recovering 的加入进一步避免用户刚回到任务时被继续打扰。"
    )

    add_heading(doc, "5.2 证据融合与权重", 2)
    add_heading(doc, "5.2.1 Normal 证据", 3)
    add_para(
        doc,
        "Normal 的证据更关注行为是否稳定，而不是是否高速阅读。"
        "状态机中，gaze 模式下面部存在可增加约 0.4，交互活跃增加约 0.3，滚动速度处于 0 到 2.0 的平滑区间增加约 0.2，"
        "dwellTime 在 1 到 6 秒范围内增加约 0.1，鼠标阅读横向比例大于 0.6 也会增加约 0.1。"
        "这些证据共同表达了“用户仍在阅读任务中且节奏可解释”。"
    )
    add_heading(doc, "5.2.2 Distracted 证据", 3)
    add_para(
        doc,
        "Distracted 的证据来自缺失和分散。gaze 模式下，如果面部缺失超过 2000ms，概率会显著上升；"
        "交互不活跃、idleDuration 超过 3000ms、mouse 模式下超过 5000ms 的长空闲，以及 heatmap dispersion 较高，"
        "都会增强分心判断。这样做可以覆盖摄像头模式和非摄像头模式下的不同分心表现。"
    )
    add_heading(doc, "5.2.3 Struggling 与 Recovering 证据", 3)
    add_para(
        doc,
        "Struggling 的核心证据是长时间停留和推进停滞。默认 strugglingDwellTime 为 8000ms，"
        "若用户 dwellTime 超过该值，会产生较强困难证据；若同时存在 scrollIdleDuration 超过 5000ms、scrollVelocity 低于 0.5、"
        "方向变化较多但缺少阅读模式，Struggling 概率会继续提高。"
    )
    add_para(
        doc,
        "Recovering 的证据体现“恢复动作”：面部重新出现、交互恢复、滚动恢复、以及当前已处于 Recovering。"
        "这使系统能识别用户从异常状态返回正常任务的过程，并为低强度反馈留下空间。"
    )

    add_table(
        doc,
        ["状态", "主要证据", "默认阈值或权重", "适合的界面响应"],
        [
            ["Normal", "面部存在、交互活跃、滚动平稳、停留适中、鼠标横向阅读轨迹。", "face +0.4；interaction +0.3；smooth scroll +0.2；dwell 1-6s +0.1。", "维持阅读环境，不主动打扰。"],
            ["Distracted", "面部离开、长时间无交互、鼠标或滚动空闲、注意点分散。", "face absent >2000ms；idle >3000ms；mouse idle >5000ms；dispersion >200。", "轻提醒、浮动提示、必要时声音或醒目提示。"],
            ["Struggling", "长时间停留、滚动暂停、低滚动速度、方向变化异常。", "dwell >8000ms；scroll idle >5000ms；velocity <0.5；direction changes >5。", "关键词高亮、摘要面板、段落解释。"],
            ["Recovering", "面部返回、交互恢复、滚动恢复。", "face returns +0.4；interaction resumes +0.3；scroll resumes +0.2。", "进度提示、积极反馈、降低提醒强度。"],
        ],
        [1.1, 2.25, 2.0, 1.4],
    )

    add_heading(doc, "5.3 状态转换机制", 2)
    add_heading(doc, "5.3.1 冷却与滞后", 3)
    add_para(
        doc,
        "状态机没有在每一次特征变化时立即转换状态，而是使用 normalTransitionCooldown、hysteresis 和概率门槛。"
        "默认 normalTransitionCooldown 为 2000ms，状态概率需要超过约 0.4 才具备转换条件，当前状态还会获得 1.3 倍的调整权重。"
        "这种滞后机制可以减少状态在 Normal 与 Distracted、Normal 与 Struggling 之间快速抖动。"
    )
    add_para(
        doc,
        "对阅读系统而言，抖动会直接破坏体验：界面如果频繁变暗、弹出、取消高亮，用户会把辅助系统看成新的干扰源。"
        "因此，状态转换机制的稳定性和视觉层的低打扰性同样重要。"
    )

    add_heading(doc, "5.3.2 mouse 模式与 gaze 模式的差异", 3)
    add_para(
        doc,
        "mouse 模式下无法可靠判断 facePresent，因此状态机更依赖交互空闲、滚动暂停、鼠标轨迹和停留时间。"
        "gaze 模式下则可以使用 faceAbsentDuration、凝视区域和校准后的 gaze point。"
        "FocusFlow 把 trackingMode 写入 config，并在 PerceptionModule 和 main.js 中共同使用，避免把摄像头相关判断错误套用到鼠标模式。"
    )
    add_para(
        doc,
        "这种差异处理也体现了渐进增强原则：基础体验不依赖高权限传感器，用户授权后再增加更精细的眼动判断。"
        "在人机交互课程项目中，这比强制摄像头更容易被用户接受，也更适合实际课堂演示。"
    )

    add_heading(doc, "5.4 状态事件与界面同步", 2)
    add_para(
        doc,
        "StateMachine 在状态变化时派发 focusflow-state-change 事件。"
        "main.js 监听该事件后更新界面状态、记录 analytics，并触发相应的视觉反馈。"
        "事件式同步让认知模块不需要了解具体 DOM 结构，也让 UI 模块可以在不修改状态机逻辑的情况下调整表现形式。"
    )
    add_para(
        doc,
        "项目还使用 focusflow-tracking-mode、focusflow-api-error、focusflow-lang-change 等事件处理模式切换、API 错误和语言变化。"
        "这种事件体系是前端项目保持模块边界的重要方式，也是后续继续拆分 main.js 的基础。"
    )


def add_section_6(doc: Document):
    add_heading(doc, "6 决策模块与自适应干预", 1)
    add_heading(doc, "6.1 DecisionModule 的职责", 2)
    add_para(
        doc,
        "js/decision/decisionModule.js 接收当前状态、features 和 userProfile，先更新 AdaptiveThreshold，"
        "再由 InterventionStrategy 根据状态和阈值选择策略，最后返回包含 strategy、thresholds、state、profile、timestamp 和 escalationLevel 的决策对象。"
        "这个模块把“判断用户状态”和“决定界面怎么做”分开，使状态机不直接承担界面策略责任。"
    )
    add_para(
        doc,
        "DecisionModule 还记录策略效果。它会根据后续状态是否改善来估计某类干预是否有效。"
        "虽然当前实现仍然是启发式效果跟踪，但它已经为个性化策略选择留下接口：如果某位用户对浮动提示反应较差，"
        "系统未来可以降低该策略权重，改用关键词或进度反馈。"
    )

    add_heading(doc, "6.2 AdaptiveThreshold 的个性化逻辑", 2)
    add_heading(doc, "6.2.1 用户统计与画像", 3)
    add_para(
        doc,
        "js/decision/adaptiveThreshold.js 保存 readingSpeedSamples、dwellTimeSamples、distractionDurations、"
        "recoveryTimes 和 interactionPatterns 等用户统计。默认用户画像中 readingSpeed 平均值约为 150，"
        "dwellTime 平均值约为 4000ms，distractionSensitivity 为 0.5，strugglingThreshold 为 8000ms。"
        "这些默认值为新用户提供初始行为边界。"
    )
    add_para(
        doc,
        "模块设置 learningRate 为 0.1，最少样本量为 10，适配间隔约 60000ms。"
        "这意味着系统不会因为几次偶然行为立即改变阈值，而是在一段阅读会话后逐步适配。"
        "对课程项目而言，这是一种简单但合理的个性化策略，既能体现自适应思想，也不会引入难以解释的复杂模型。"
    )

    add_heading(doc, "6.2.2 阈值调整原则", 3)
    add_para(
        doc,
        "AdaptiveThreshold 会根据阅读速度调整 struggling dwell time：慢读用户获得更长容忍时间，快读用户的困难阈值相对缩短。"
        "它还会根据个人平均 dwellTime 形成至少 4000ms 的个人阈值，并根据分心频率调整 distraction sensitivity。"
        "如果用户恢复速度较快，recoveringStableTime 可以缩短，使系统更快回到 Normal。"
    )
    add_para(
        doc,
        "这种策略尊重阅读差异。慢读不应被系统误判为困难，快速跳读也不应被强迫放慢。"
        "自适应阈值的价值不在于绝对准确，而在于让系统随着用户行为积累逐渐减少固定阈值带来的误判。"
    )

    add_heading(doc, "6.3 干预策略矩阵", 2)
    add_table(
        doc,
        ["状态", "策略 ID", "界面形式", "强度 / 时长 / 冷却", "设计理由"],
        [
            ["Distracted", "SUBTLE_OVERLAY", "轻微覆盖层或环境变化。", "0.3 / 3000ms / 15000ms", "先用低强度方式唤回用户，避免一开始就强提醒。"],
            ["Distracted", "FLOATING_PROMPT", "浮动提示。", "0.5 / 4000ms / 20000ms", "在轻提示无效时提供更明确的任务回归信号。"],
            ["Distracted", "SOUND_ALERT", "声音或强提醒。", "0.8 / 2000ms / 30000ms", "仅在升级情况下使用，冷却较长，避免频繁打扰。"],
            ["Struggling", "KEYWORD_HIGHLIGHT", "关键词高亮。", "0.3 / 5000ms / 20000ms", "先帮助用户抓住段落结构，不改变阅读路径。"],
            ["Struggling", "SUMMARY_PANEL", "摘要面板或理解卡片。", "0.6 / 10000ms / 30000ms", "用户停滞较明显时提供内容压缩和解释。"],
            ["Struggling", "SIMPLIFICATION", "简化或更强辅助。", "0.8 / 15000ms / 60000ms", "作为高强度帮助，冷却最长，避免替代阅读。"],
            ["Recovering", "PROGRESS_INDICATOR", "阅读进度提示。", "0.2 / 3000ms / 10000ms", "帮助用户确认已回到任务。"],
            ["Recovering", "POSITIVE_FEEDBACK", "积极反馈。", "0.3 / 3000ms / 20000ms", "强化恢复行为而不制造压力。"],
        ],
        [0.95, 1.35, 1.45, 1.35, 1.65],
        font_size=8.4,
    )

    add_heading(doc, "6.4 低打扰交互原则", 2)
    add_para(
        doc,
        "InterventionStrategy 对同一状态连续出现进行 escalation level 判断，并计算 severity。"
        "severity 综合状态置信度、状态持续时间、升级等级以及状态相关 features。"
        "这说明项目不是只根据状态名称选择策略，而是考虑状态强度和历史上下文。"
    )
    add_para(
        doc,
        "策略执行还受到冷却时间约束。每种策略有独立 cooldown，策略之间还有全局最小间隔。"
        "这对用户体验非常重要，因为阅读任务需要连续心流，系统提醒越多，越可能成为新的干扰。"
        "FocusFlow 的策略矩阵从轻到重递进，符合“先不打断，再适当帮助，最后才强提醒”的交互原则。"
    )
    add_para(
        doc,
        "在课程答辩中，决策模块可以作为项目创新点说明：它不只是检测分心后弹窗，而是根据分心、困难和恢复三类非正常状态选择不同干预。"
        "这种区分使系统更接近真实自适应界面，而不是普通计时提醒工具。"
    )


def add_section_7(doc: Document):
    add_heading(doc, "7 阅读界面、内容处理与反馈层", 1)
    add_heading(doc, "7.1 ReadingView 的阅读块模型", 2)
    add_heading(doc, "7.1.1 文本加载与块元素映射", 3)
    add_para(
        doc,
        "js/ui/readingView.js 是前端阅读体验的主模块。它负责加载默认文本或用户导入文本，识别文档语言，"
        "构建 blockElements，维护当前阅读块，更新 WPM 和理解辅助按钮。"
        "这种阅读块模型是后续感知和辅助的基础：没有块映射，系统只能知道用户在页面某处停留，却无法知道停留对应哪段内容。"
    )
    add_para(
        doc,
        "ReadingView 还与 main.js 中的 gazeBlock 计算相连。"
        "当用户凝视或指针落在某个块上，main.js 可以调用 readingView.getBlockAtGaze(data.x, data.y)，"
        "然后将 block index、block text、reading progress 和 WPM 传给摘要、分析和视觉反馈模块。"
        "这使项目形成了从坐标到文本语义的桥接。"
    )

    add_heading(doc, "7.1.2 段落辅助按钮", 3)
    add_para(
        doc,
        "ReadingView 为阅读块维护理解辅助按钮状态，例如 hidden、generate、loading、reopen 等。"
        "main.js 在 block summary cache 存在时把按钮更新为 reopen，在摘要生成中把按钮隐藏或显示加载状态。"
        "这种状态机式按钮管理使用户可以主动打开摘要，而不是被系统强制展示全部解释。"
    )
    add_para(
        doc,
        "从 HCI 角度看，主动触发与系统推荐的结合很重要。"
        "当系统判断用户可能 struggling 时，可以把摘要入口呈现出来；但是否阅读摘要仍由用户决定。"
        "这保留了用户控制权，避免辅助功能变成替代性学习。"
    )

    add_heading(doc, "7.2 ParagraphSplitter 与文本导入", 2)
    add_para(
        doc,
        "js/ui/paragraphSplitter.js 负责把原始文本解析为可阅读块。"
        "它会处理换行、段落合并、长段落拆分、标题识别和标题块标记。"
        "长文本导入后，如果不进行合理拆分，阅读块过大就无法准确定位用户停留位置；块过小又会让界面碎片化。"
        "该模块在两者之间取得平衡，使阅读器更适合课程材料和文章类文本。"
    )
    add_para(
        doc,
        "js/utils/textEncoding.js 则处理文件编码识别和文本解码。"
        "它通过尝试不同编码并评分文本质量来减少中文或英文材料导入时的乱码风险。"
        "这类工具代码在界面上不显眼，但对课程演示非常关键，因为导入文件失败会直接影响系统可用性。"
    )

    add_heading(doc, "7.3 VisualEffects 的反馈实现", 2)
    add_heading(doc, "7.3.1 视觉效果类型", 3)
    add_para(
        doc,
        "js/ui/visualEffects.js 是项目中体量较大的 UI 文件，负责 dim overlay、gaze glow、关键词高亮、呼吸效果、浮动提示、"
        "wakeup overlay、warm flash、进度提示和 comprehension card 等视觉反馈。"
        "这些效果对应决策层中的干预策略，使系统能够用不同强度和形式回应不同状态。"
    )
    add_para(
        doc,
        "视觉层需要特别注意“帮助”和“打扰”的边界。"
        "FocusFlow 使用覆盖层、轻微高亮、卡片和进度提示等方式，而不是将用户强制跳转到另一个页面。"
        "这种设计让反馈停留在阅读上下文内部，减少用户重新定位阅读位置的成本。"
    )

    add_heading(doc, "7.3.2 焦点模式与界面降噪", 3)
    add_para(
        doc,
        "js/ui/focusMode.js 管理专注模式及相关键盘快捷行为。"
        "专注模式的意义是减少阅读页面上无关元素的视觉竞争，使用户更容易维持当前任务。"
        "它与状态干预不同：状态干预是在系统检测到问题时响应，专注模式则是用户主动选择的环境调整。"
    )
    add_para(
        doc,
        "js/ui/debugPanel.js 维护 FPS、gaze/block/state/dim/focus mode/WPM/effects count 等调试信息。"
        "调试面板不属于最终用户的主要阅读界面，但对课程项目开发和答辩很有价值，因为它能展示系统内部状态不是凭空描述，"
        "而是来自实时特征和状态更新。"
    )

    add_heading(doc, "7.4 多语言界面与可访问性", 2)
    add_para(
        doc,
        "js/i18n/i18n.js 包含大量界面翻译和语言切换逻辑，并通过 focusflow-lang-change 事件更新页面文本。"
        "这使系统能够服务中文和英文材料，避免阅读辅助只适用于单一语言。"
        "对于 HCI 作业，多语言不是装饰性功能，而是影响真实用户理解和可用性的基础能力。"
    )
    add_para(
        doc,
        "可访问性方面，系统仍有提升空间。例如部分视觉效果需要保证对低视力用户足够清晰，声音提醒需要提供开关，"
        "键盘操作和屏幕阅读器语义也可以进一步完善。"
        "不过当前模块划分已经为这些改进提供了入口：视觉效果集中在 VisualEffects，文本集中在 i18n，阅读结构集中在 ReadingView。"
    )


def add_section_8(doc: Document):
    add_heading(doc, "8 NLP 辅助与摘要服务", 1)
    add_heading(doc, "8.1 KeywordExtractor 的关键词逻辑", 2)
    add_para(
        doc,
        "js/nlp/keywordExtractor.js 使用分词、停用词过滤、词频、双词组合和块级提取等方法寻找段落关键词。"
        "关键词高亮与 Struggling 状态高度相关：当用户在段落中停留过久时，系统不必立即给出完整答案，"
        "可以先突出概念词、主题词或关键短语，帮助用户重新组织段落结构。"
    )
    add_para(
        doc,
        "关键词提取属于可解释辅助。与直接给出大段总结相比，关键词更像阅读支架，"
        "它提醒用户哪些信息值得注意，但仍保留用户自主理解文本的空间。"
        "因此它适合作为低强度 struggling 干预。"
    )

    add_heading(doc, "8.2 ParagraphSummarizer 的规则摘要", 2)
    add_para(
        doc,
        "js/nlp/paragraphSummarizer.js 提供规则摘要能力，支持语言检测、中文和英文句子切分、关键词、主题句、细节句和压缩比例控制。"
        "它的存在使系统在没有 LLM 服务时仍能提供基本理解辅助。"
        "这对课程演示很重要，因为外部 API 可能因网络、密钥或费用限制不可用。"
    )
    add_para(
        doc,
        "规则摘要的质量通常不如大型语言模型，但它具有可控、快速和离线友好的优势。"
        "在 FocusFlow 的系统结构中，规则摘要不是临时替代，而是 fallback 机制的一部分。"
        "这种降级路径提高了系统鲁棒性。"
    )

    add_heading(doc, "8.3 LLMSummaryManager 与服务端代理", 2)
    add_heading(doc, "8.3.1 缓存、预取与队列", 3)
    add_para(
        doc,
        "js/nlp/llmSummaryManager.js 管理 LLM 摘要请求、缓存、预取、并发队列和文档语言。"
        "main.js 在用户稳定阅读某个 block 时调用 onReadingBlock，让摘要在用户仍阅读时提前准备。"
        "当用户真正需要理解辅助时，系统可以更快显示结果，减少等待带来的阅读中断。"
    )
    add_para(
        doc,
        "配置中 llmConcurrency 默认为 3，llmApiUrl 为 /api/summarize，llmStatusUrl 为 /api/llm/status。"
        "这些配置集中在 main.js 的 config 中，便于替换后端接口或关闭 LLM。"
        "摘要缓存还避免同一段落重复请求，减少响应延迟和外部服务成本。"
    )

    add_heading(doc, "8.3.2 服务端 prompt 与错误处理", 3)
    add_para(
        doc,
        "server/dev-server.js 的 buildPrompt(text, lang) 会根据 zh/en 语言选择摘要提示，callLLM 调用 chat/completions 接口。"
        "当 API key 未配置时，服务返回 NO_API_KEY；当上游响应失败时，返回 LLM_HTTP_ERROR 或 LLM_ERROR。"
        "前端可以据此显示 API 错误事件 focusflow-api-error，或者进入规则摘要 fallback。"
    )
    add_para(
        doc,
        "这种服务端代理设计比在浏览器端直接暴露 API key 更安全。"
        "它也把外部服务调用从 UI 层移出，使前端模块只关心摘要是否可用，而不用处理密钥和 HTTP 细节。"
    )

    add_heading(doc, "8.4 摘要辅助的交互边界", 2)
    add_para(
        doc,
        "理解辅助不应让用户完全跳过原文。"
        "FocusFlow 把摘要入口放在段落附近，并通过按钮状态控制显示时机，使摘要成为用户遇到困难时的支架。"
        "对于课程项目来说，这种边界比单纯展示 AI 摘要更有 HCI 意义：系统关注的是人在阅读中的状态，而不是只展示模型能力。"
    )
    add_para(
        doc,
        "后续可以进一步把摘要效果纳入 analytics。例如记录用户打开摘要后是否恢复滚动、是否减少 dwellTime、是否从 Struggling 转为 Recovering。"
        "这样可以评估摘要是否真的帮助阅读，而不是只统计按钮点击次数。"
    )


def add_section_9(doc: Document):
    add_heading(doc, "9 数据分析、会话报告与评估方案", 1)
    add_heading(doc, "9.1 AttentionAnalytics 的指标体系", 2)
    add_para(
        doc,
        "js/analytics/attentionAnalytics.js 记录 gaze sample、分心开始和结束、阅读块词数、状态转换、理解辅助使用和自动保存历史。"
        "它提供 getAttentionScore、getReadingSpeed、getSessionDuration、getFocusRatio、getRegressionRate、"
        "getAverageRecoverySeconds、getBlockHeatmap 等接口。"
        "这些指标把系统内部事件转化为可以解释的学习过程数据。"
    )
    add_para(
        doc,
        "Attention score 不能被理解为对人的能力评价，而应理解为一次阅读会话中的行为稳定性指标。"
        "Focus ratio 可以反映 Normal 或有效阅读状态所占比例，recovery seconds 可以反映干预后回到任务的速度，"
        "block heatmap 可以显示哪些段落消耗更多注意力。"
        "这些指标适合用于项目评估和个人反思，但不适合作为高风险评价。"
    )

    add_heading(doc, "9.2 SessionReport 的可视化输出", 2)
    add_para(
        doc,
        "js/analytics/sessionReport.js 提供 show、_statCard、_renderHeatmap、_renderTimeline、_exportJson 等功能。"
        "它可以把会话持续时间、阅读速度、专注比例、热力图和状态时间线集中显示在 ff-session-report-overlay 中。"
        "这使课程演示可以展示“系统做了什么”和“用户状态怎样变化”，而不是只展示阅读页面。"
    )
    add_para(
        doc,
        "会话报告对迭代也有帮助。"
        "如果报告显示频繁误判 Struggling，说明 dwellTime 或滚动暂停阈值需要调整；"
        "如果报告显示分心提醒过多但恢复速度没有提升，说明干预策略可能过强或时机不对。"
        "因此，分析模块不仅用于结果展示，也用于诊断系统设计。"
    )

    add_heading(doc, "9.3 课程评估方案", 2)
    add_heading(doc, "9.3.1 实验任务设计", 3)
    add_para(
        doc,
        "评估可以采用同一组长文本材料，让参与者分别使用普通阅读界面和 FocusFlow 阅读界面完成阅读任务。"
        "材料应包含事实信息、概念解释和一段相对复杂的论证内容，以便观察分心、困难和恢复三类状态。"
        "每次阅读后设置理解题和主观体验问卷，同时记录系统会话指标。"
    )
    add_para(
        doc,
        "为了避免学习效应，实验可以使用 A/B 顺序平衡：一半参与者先用普通界面再用 FocusFlow，另一半反过来。"
        "材料难度尽量接近，阅读时间控制在可承受范围内。"
        "如果使用 gaze 模式，应在正式任务前完成九点校准；如果摄像头条件不稳定，则统一使用 mouse 模式保证数据可比。"
    )

    add_heading(doc, "9.3.2 指标设计", 3)
    add_table(
        doc,
        ["指标类型", "指标", "采集方式", "评价意义"],
        [
            ["客观行为", "阅读完成时间", "Session duration 与任务计时", "观察辅助是否显著拖慢或改善阅读效率。"],
            ["客观行为", "专注比例", "AttentionAnalytics.getFocusRatio()", "衡量有效阅读状态所占比例。"],
            ["客观行为", "平均恢复时间", "getAverageRecoverySeconds()", "判断提示是否帮助用户更快回到任务。"],
            ["客观行为", "理解辅助使用次数", "recordComprehensionAssist()", "观察摘要和高亮是否被用户真实使用。"],
            ["客观行为", "状态转换次数", "recordStateTransition()", "反映状态稳定性和误判风险。"],
            ["学习结果", "理解题正确率", "阅读后测验", "评估辅助是否改善理解，而不仅是改变界面行为。"],
            ["主观体验", "打扰感评分", "Likert 量表", "评估提示强度是否过高。"],
            ["主观体验", "控制感与信任", "访谈或问卷", "判断用户是否理解并接受系统反馈。"],
        ],
        [1.05, 1.55, 2.0, 2.1],
    )

    add_heading(doc, "9.4 数据解释原则", 2)
    add_para(
        doc,
        "评估结果不能只看阅读速度。阅读速度变快可能来自跳读，也可能来自理解更顺畅；阅读速度变慢可能代表深入阅读，也可能代表系统干扰。"
        "因此需要把速度、理解题、专注比例、恢复时间和主观体验一起解释。"
        "FocusFlow 的指标体系适合做多维分析，而不是用一个分数代表全部体验。"
    )
    add_para(
        doc,
        "另外，系统记录的 gaze 或 mouse 数据属于行为数据，应只用于当前学习辅助和项目评估。"
        "评估时应告知参与者采集内容、使用目的和退出方式。"
        "这不仅是伦理要求，也能提高参与者对系统的信任。"
    )


def add_section_10(doc: Document):
    add_heading(doc, "10 代码质量、风险与功能优化方向", 1)
    add_heading(doc, "10.1 已完成的功能优化", 2)
    add_para(
        doc,
        "当前项目已经从单一眼动演示扩展为更完整的阅读辅助系统。"
        "默认 mouse tracking 提高了可运行性；gaze 模式保留 WebGazer 与校准能力；"
        "KalmanFilter2D 改善凝视坐标稳定性；LLMSummaryManager 增加摘要缓存和预取；"
        "SessionReport 让会话数据可以被展示和导出；cameraAccess 与 textEncoding 增强了权限和文件导入的容错。"
    )
    add_para(
        doc,
        "这些优化共同解决了课程项目最常见的风险：摄像头不可用、眼动抖动、导入文本乱码、摘要服务失败、状态变化无法解释、演示后缺少结果数据。"
        "项目现在不再依赖某一个单点功能，而是形成了可降级的功能链。"
    )

    add_heading(doc, "10.2 主要代码风险", 2)
    add_heading(doc, "10.2.1 main.js 职责过重", 3)
    add_para(
        doc,
        "main.js 同时处理初始化、模式切换、校准、感知更新、状态决策、摘要缓存、按钮状态、会话报告和销毁。"
        "这种集中控制便于早期开发和课堂演示，但随着功能增加，文件会越来越难维护。"
        "例如摘要相关逻辑、策略执行逻辑、tracking mode 管理逻辑都可以进一步拆成专门模块。"
    )
    add_para(
        doc,
        "优化方向是保留 main.js 的编排角色，把具体职责外移：TrackingController 管理 mouse/gaze 切换和校准；"
        "AssistanceController 管理摘要、按钮和理解卡片；InterventionExecutor 管理策略到视觉效果的映射。"
        "这样可以减少单文件修改引发的回归风险。"
    )

    add_heading(doc, "10.2.2 视觉效果与策略映射需要更清晰", 3)
    add_para(
        doc,
        "InterventionStrategy 给出策略对象，VisualEffects 实现实际效果，但两者之间的映射部分仍依赖 main.js 中的过程逻辑。"
        "如果后续增加更多策略，容易出现策略已定义但 UI 没有对应实现，或 UI 效果存在但决策层不会触发的问题。"
        "更稳妥的方式是建立策略 ID 到执行函数的注册表，并为每个策略声明进入条件、退出条件和视觉资源。"
    )
    add_para(
        doc,
        "这种改造可以提高可测试性。"
        "开发者可以单独测试 SUBTLE_OVERLAY、SUMMARY_PANEL、PROGRESS_INDICATOR 等策略是否在给定状态下被正确执行，"
        "而不必完整启动摄像头和 WebGazer。"
    )

    add_heading(doc, "10.2.3 状态阈值仍然依赖启发式", 3)
    add_para(
        doc,
        "当前状态判断以启发式权重为主，优点是可解释，缺点是跨用户稳定性有限。"
        "不同用户的阅读习惯、鼠标使用习惯、屏幕大小和文本难度都会影响 dwellTime 与 scrollVelocity。"
        "AdaptiveThreshold 已经提供个性化方向，但仍需要更多真实数据验证阈值是否合适。"
    )
    add_para(
        doc,
        "后续可以在保留可解释规则的基础上引入轻量分类器，例如用规则输出作为特征，再通过少量标注数据训练逻辑回归或决策树。"
        "这样既能提高准确率，也不会完全牺牲课程项目所需的可解释性。"
    )

    add_heading(doc, "10.3 功能优化建议", 2)
    add_table(
        doc,
        ["优化方向", "当前基础", "建议做法", "预期收益"],
        [
            ["模块拆分", "main.js 负责过多流程。", "拆出 TrackingController、AssistanceController、InterventionExecutor。", "降低维护难度，减少功能之间相互影响。"],
            ["策略注册", "策略和 UI 效果存在映射但不够显式。", "建立 strategyId → handler 的注册表，并统一冷却和退出逻辑。", "增加策略时更稳定，便于测试。"],
            ["阈值验证", "AdaptiveThreshold 提供个性化但缺少实验数据。", "用小规模用户实验校准 dwell、idle、scroll 相关阈值。", "降低误判，提升用户信任。"],
            ["可访问性", "已有多语言和焦点模式。", "补充键盘导航、ARIA 标签、对比度检查、声音提醒开关。", "让辅助系统适配更多用户。"],
            ["隐私控制", "默认 mouse 模式和摄像头工具封装已完成。", "加入数据采集提示、会话清除按钮、本地保存范围说明。", "减少用户对监控感的担忧。"],
            ["评估闭环", "SessionReport 已能展示会话数据。", "把干预前后状态变化与理解题结果关联。", "证明辅助不仅改变界面，还改善阅读效果。"],
            ["错误恢复", "LLM 和摄像头有回退路径。", "在界面上更清晰地区分 disabled、loading、fallback、error。", "提升演示可靠性和用户理解。"],
            ["国际化质量", "i18n 文件覆盖大量界面文本。", "对中英文长文本按钮、状态短语和摘要提示做统一术语表。", "减少界面语言不一致。"],
        ],
        [1.05, 1.7, 2.05, 1.85],
    )

    add_heading(doc, "10.4 与课程要求的对应关系", 2)
    add_para(
        doc,
        "人机交互课程通常关注用户、任务、界面、反馈、评估和伦理。"
        "FocusFlow 对应这些要求的方式比较清晰：用户是长文本学习者，任务是持续阅读和理解，界面是可自适应阅读器，"
        "反馈是不同强度的视觉与内容辅助，评估来自会话指标与理解测试，伦理重点是摄像头权限和行为数据使用边界。"
    )
    add_para(
        doc,
        "项目最适合作为课程作业的原因是它能够把 HCI 理论和可运行代码连接起来。"
        "状态机、阈值、冷却时间、视觉反馈、摘要入口、会话报告都能在代码中找到对应实现，"
        "答辩时可以从界面演示切换到源码解释，证明设计不是停留在概念层面。"
    )


def add_section_11(doc: Document):
    add_heading(doc, "11 代码文件索引与实现细节", 1)
    add_heading(doc, "11.1 核心文件索引", 2)
    add_table(
        doc,
        ["文件", "主要类或函数", "关键职责", "值得关注的实现点"],
        [
            ["index.html", "initGazeDot、initTheme、timer 相关函数", "承载阅读界面、控制区、状态展示和主题逻辑。", "DOM 节点命名统一使用 ff-*，便于模块选择与事件绑定。"],
            ["js/main.js", "FocusFlow 对象及初始化/更新函数", "全局编排感知、认知、决策、UI、NLP 与分析模块。", "默认 mouse 模式，gaze 模式经 WebGazer 与校准进入；负责摘要预取和会话报告。"],
            ["js/calibration/calibrationManager.js", "CalibrationManager", "九点校准覆盖层、样本收集、跳过与完成回调。", "以 3×3 屏幕点提升 WebGazer 个体化准确度。"],
            ["js/perception/perceptionModule.js", "PerceptionModule", "整合面部、头部、凝视、鼠标、滚动和热力图特征。", "getFeatures() 是认知层输入的统一出口。"],
            ["js/perception/mouseTracker.js", "MouseTracker", "记录鼠标速度、轨迹、方向变化和阅读/搜索模式。", "支持无摄像头情况下的注意力近似判断。"],
            ["js/perception/scrollAnalyzer.js", "ScrollAnalyzer", "分析滚动速度、暂停、空闲、爆发和阅读节奏。", "与 dwellTime 共同判断 struggling。"],
            ["js/perception/gazeRegion.js", "GazeRegionMapping", "映射页面区域、维护 dwell 和 gaze trail。", "把坐标转化为阅读区域和停留时间。"],
            ["js/perception/kalmanFilter.js", "KalmanFilter2D", "平滑二维凝视坐标。", "降低 WebGazer 抖动对高亮和块定位的影响。"],
            ["js/cognition/stateMachine.js", "StateMachine", "四状态概率判断、转换冷却、滞后和事件派发。", "权重可解释，支持 mouse/gaze 差异化证据。"],
            ["js/decision/decisionModule.js", "DecisionModule", "连接阈值适配与策略选择。", "返回完整 decision 对象，记录策略效果。"],
            ["js/decision/adaptiveThreshold.js", "AdaptiveThreshold", "根据用户阅读速度、停留和恢复数据调整阈值。", "learningRate、min samples 和 adaptation interval 控制适配稳定性。"],
            ["js/decision/interventionStrategy.js", "InterventionStrategy", "定义 distraction、struggling、recovery 干预策略。", "含强度、时长、冷却、升级和 severity 计算。"],
            ["js/ui/readingView.js", "ReadingView", "文本加载、阅读块、WPM、当前块和理解按钮。", "把页面位置与段落内容关联。"],
            ["js/ui/paragraphSplitter.js", "ParagraphSplitter", "段落拆分、标题识别和长段落切分。", "影响阅读块粒度和后续分析准确性。"],
            ["js/ui/visualEffects.js", "VisualEffects", "覆盖层、高亮、光标、提示、理解卡片和恢复反馈。", "实现决策策略在界面上的可见形式。"],
            ["js/ui/focusMode.js", "FocusMode", "专注模式与键盘操作。", "支持用户主动降噪。"],
            ["js/ui/debugPanel.js", "DebugPanel", "调试状态、FPS、WPM、效果数量和状态显示。", "为开发与答辩提供内部状态证据。"],
            ["js/nlp/keywordExtractor.js", "KeywordExtractor", "关键词、双词组合和块级关键词。", "支撑 KEYWORD_HIGHLIGHT。"],
            ["js/nlp/paragraphSummarizer.js", "ParagraphSummarizer", "规则摘要与语言处理。", "作为 LLM 不可用时的本地回退。"],
            ["js/nlp/llmSummaryManager.js", "LLMSummaryManager", "摘要缓存、预取、队列和语言设置。", "减少摘要等待并降低重复请求。"],
            ["js/analytics/attentionAnalytics.js", "AttentionAnalytics", "状态、速度、热力图、理解辅助和历史记录。", "提供项目评估指标。"],
            ["js/analytics/sessionReport.js", "SessionReport", "统计卡片、热力图、时间线和 JSON 导出。", "把会话数据转化为可展示结果。"],
            ["js/utils/cameraAccess.js", "CameraAccess", "摄像头权限、支持检测、错误分类和释放。", "提高 gaze 模式的可靠性和隐私可控性。"],
            ["js/utils/textEncoding.js", "TextEncoding", "文本编码检测、解码和质量评分。", "降低中文材料导入乱码风险。"],
            ["server/dev-server.js", "loadEnvFile、callLLM、serveStatic", "静态服务和 LLM 摘要代理。", "避免在浏览器端暴露 API key，并提供 fallback 条件。"],
        ],
        [1.45, 1.45, 2.0, 1.85],
        font_size=7.8,
    )

    add_heading(doc, "11.2 关键事件与 DOM 节点", 2)
    add_para(
        doc,
        "项目通过事件和 DOM id 将多个模块连接起来。"
        "关键事件包括 focusflow-state-change、focusflow-tracking-mode、focusflow-api-error、focusflow-lang-change、DOMContentLoaded、beforeunload、click、mousemove 和 scroll。"
        "关键节点包括 ff-reading-content、ff-reading-panel、ff-camera-gate、ff-camera-gate-btn、ff-tracking-toggle、ff-gaze-cursor、"
        "ff-dim-overlay、ff-wakeup-overlay、ff-comprehension-card、ff-session-report-overlay 等。"
    )
    add_para(
        doc,
        "这些命名体现了项目的前端组织方式：模块不通过全局随机选择器互相寻找，而是围绕统一前缀和自定义事件协作。"
        "后续若进行组件化重构，可以把这些节点封装为独立组件，但现阶段的命名已经足以支持课程演示和代码阅读。"
    )

    add_heading(doc, "11.3 状态到反馈的实现链", 2)
    add_para(
        doc,
        "以用户长时间停留在某段为例，ScrollAnalyzer 和 GazeRegionMapping 会产生 dwellTime 与 scrollIdleDuration，"
        "PerceptionModule 将其合并为 features，StateMachine 在 Struggling 证据增强后派发状态变化，"
        "DecisionModule 选择 KEYWORD_HIGHLIGHT 或 SUMMARY_PANEL，VisualEffects 与 ReadingView 在对应段落附近呈现高亮或摘要入口，"
        "AttentionAnalytics 记录该次状态转换和理解辅助行为。"
    )
    add_para(
        doc,
        "以用户离开摄像头为例，gaze 模式下 FaceDetection 的 faceAbsentDuration 持续上升，"
        "StateMachine 增强 Distracted 概率，InterventionStrategy 先选择 SUBTLE_OVERLAY，若状态持续再升级到 FLOATING_PROMPT。"
        "如果用户返回并恢复交互，Recovering 证据增强，系统转向进度提示或积极反馈，而不是继续强提醒。"
    )

    add_heading(doc, "11.4 核心函数链路", 2)
    add_para(
        doc,
        "源码中最值得关注的不是单个函数名称，而是函数之间怎样形成闭环。"
        "FocusFlow 的前端代码大量使用对象方法和浏览器事件，典型链路从初始化开始，经过行为采样、特征生成、状态判断、策略选择、界面执行和数据记录。"
        "这些链路如果只看界面很难发现，但通过代码可以清楚看到每一层的责任边界。"
    )
    add_table(
        doc,
        ["链路阶段", "关键入口", "输入", "输出", "代码分析"],
        [
            ["应用启动", "FocusFlow.init()", "页面 DOM、config、默认文本、服务状态。", "各模块实例与默认 mouse tracking。", "初始化顺序先 UI 后感知，再认知、决策、视觉、调试、NLP 和分析，保证后续事件处理时依赖对象已存在。"],
            ["鼠标模式", "FocusFlow.startMouseTracking()", "mousemove、click、scroll、wheel。", "鼠标位置、滚动状态、交互活跃度。", "默认使用 mouse 模式能绕开摄像头权限，是课程演示稳定性的关键。"],
            ["眼动模式", "startGazeTracking / WebGazer listener", "摄像头授权、WebGazer prediction、校准数据。", "gaze x/y、face 状态、gaze block。", "gaze 模式是渐进增强，失败后可回到 mouse 模式，不会让系统整体不可用。"],
            ["校准流程", "CalibrationManager.start()", "九点坐标、用户点击或 Space、gaze sample。", "calibrationData 与完成回调。", "全屏覆盖层将校准作为明确任务，降低用户不知道看哪里的风险。"],
            ["特征获取", "PerceptionModule.getFeatures()", "各子模块当前状态。", "features 对象。", "把复杂输入统一成状态机能理解的字段，是分层架构中最关键的接口。"],
            ["状态更新", "StateMachine.update()", "features 与时间戳。", "当前状态、概率、持续时间。", "证据权重、冷却和滞后共同保证状态稳定，避免界面来回闪动。"],
            ["决策选择", "DecisionModule.decide()", "state、features、userProfile。", "decision 对象与 strategy。", "阈值适配和策略选择被集中在决策层，便于后续替换策略算法。"],
            ["视觉执行", "VisualEffects 方法与 ReadingView 方法", "strategy、block index、block text。", "高亮、卡片、覆盖层、提示和按钮状态。", "反馈保留在阅读上下文中，减少跳出当前任务的成本。"],
            ["摘要辅助", "LLMSummaryManager.onReadingBlock()", "稳定阅读块文本、文档语言。", "缓存摘要、队列任务、fallback 摘要。", "预取减少等待时间，fallback 保证无 API 环境下功能不完全失效。"],
            ["会话分析", "AttentionAnalytics 与 SessionReport", "状态转换、阅读速度、辅助使用、热力点。", "指标、时间线、热力图和导出 JSON。", "让项目能够被评估，而不是只停留在界面展示。"],
        ],
        [1.0, 1.45, 1.35, 1.35, 1.6],
        font_size=7.8,
    )

    add_heading(doc, "11.5 数据字段流向", 2)
    add_para(
        doc,
        "FocusFlow 中的数据字段从底层事件逐步变得抽象。最底层是坐标、速度、滚动位置、点击时间、摄像头状态等原始数据；"
        "中间层变成 dwellTime、idleDuration、scrollVelocity、interactionActive、gazeBlock 等特征；"
        "再往上变成 Normal、Distracted、Struggling、Recovering 等认知状态；"
        "最后变成 overlay、prompt、summary、highlight、progress 等界面行为。"
        "这个抽象过程是项目区别于普通网页效果的关键。"
    )
    add_para(
        doc,
        "例如 mousemove 的原始坐标本身没有语义，只有经过 MouseTracker 计算速度、方向变化、水平移动比例后，才可能成为阅读模式证据；"
        "scrollY 本身也不能说明用户是否理解，只有结合停留时间、滚动暂停和当前阅读块，才可能表示困难或精读。"
        "状态机把这些弱证据组合起来，而决策层再把组合结果转为可执行策略。"
    )
    add_table(
        doc,
        ["字段层级", "代表字段", "生成位置", "使用位置", "设计意义"],
        [
            ["原始输入", "clientX、clientY、scrollY、WebGazer prediction、click time", "浏览器事件与 WebGazer 回调", "MouseTracker、ScrollAnalyzer、GazeRegionMapping", "保留最接近用户行为的输入，但不直接作为状态结论。"],
            ["感知特征", "dwellTime、mouseVelocity、scrollVelocity、faceAbsentDuration", "PerceptionModule 及子模块", "StateMachine", "把低层行为转成稳定、可解释、可比较的字段。"],
            ["派生特征", "interactionActive、scrollIdleDuration、dispersion、gazeBlock", "PerceptionModule、AttentionHeatmap、ReadingView", "StateMachine、LLM、Analytics", "将用户行为与页面结构和阅读块关联。"],
            ["认知状态", "Normal、Distracted、Struggling、Recovering", "StateMachine", "DecisionModule、VisualEffects、Analytics", "让系统从数据处理进入人机交互语义层。"],
            ["策略对象", "strategy id、duration、intensity、cooldown", "InterventionStrategy", "main.js 与 VisualEffects", "定义干预强度和时机，避免 UI 随意响应。"],
            ["结果指标", "attention score、focus ratio、reading speed、recovery seconds", "AttentionAnalytics", "SessionReport、评估方案", "把会话过程转化为可讨论和可改进的数据。"],
        ],
        [1.05, 1.7, 1.45, 1.45, 1.75],
        font_size=8.0,
    )
    add_para(
        doc,
        "这种字段流向也暴露了后续测试的重点：如果界面反馈异常，需要先判断是原始输入没有采集到、特征派生错误、状态判断不稳定，"
        "还是策略执行没有映射到正确的视觉效果。"
        "分层检查比盲目修改 UI 更有效。"
    )

    add_heading(doc, "11.6 典型交互场景的源码追踪", 2)
    add_para(
        doc,
        "场景一是用户正常阅读。mousemove 或 gaze prediction 更新后，PerceptionModule 得到平稳坐标和适中 dwellTime，"
        "ScrollAnalyzer 记录低到中等滚动速度，StateMachine 的 Normal 概率保持较高，DecisionModule 不触发强干预。"
        "ReadingView 继续更新当前块和 WPM，Analytics 记录正常阅读时间。"
        "这一场景体现系统的克制：正常时不刷存在感。"
    )
    add_para(
        doc,
        "场景二是用户遇到难段落。当前 block 的 dwellTime 超过 8000ms，滚动速度低于 0.5 或滚动暂停超过 5000ms，"
        "StateMachine 将 Struggling 概率推高。DecisionModule 会优先选择 KEYWORD_HIGHLIGHT 或 SUMMARY_PANEL，"
        "main.js 根据 block index 获取文本并调用关键词或摘要逻辑。"
        "如果 LLM 已预取，理解卡片可以快速显示；如果没有外部服务，则使用 ParagraphSummarizer 的规则摘要。"
    )
    add_para(
        doc,
        "场景三是用户离开阅读任务。在 mouse 模式下，长时间没有鼠标和滚动输入会提升 Distracted 概率；"
        "在 gaze 模式下，faceAbsentDuration 超过 2000ms 是更强证据。"
        "InterventionStrategy 不会马上使用最强提醒，而是先选择 SUBTLE_OVERLAY，只有连续状态和 severity 上升后才升级。"
        "这体现了低打扰和渐进式反馈原则。"
    )
    add_para(
        doc,
        "场景四是用户恢复阅读。交互恢复、滚动恢复或面部重新出现后，Recovering 的证据增强。"
        "系统在这一阶段更适合显示 PROGRESS_INDICATOR 或 POSITIVE_FEEDBACK，而不是继续分心提醒。"
        "Analytics 记录恢复时间，后续可用于评估干预是否有效。"
    )


def add_section_12(doc: Document):
    add_heading(doc, "12 运行流程、异常处理与验收要点", 1)
    add_heading(doc, "12.1 运行流程", 2)
    add_heading(doc, "12.1.1 本地启动", 3)
    add_para(
        doc,
        "项目通过 Node 本地服务运行，package.json 中的 start 和 serve 脚本都指向 server/dev-server.js。"
        "启动后，服务提供 index.html、js 目录、样式和静态资源，同时开放 LLM 状态和摘要接口。"
        "这种运行方式比直接打开 HTML 更稳定，因为摘要接口和部分浏览器权限通常需要本地服务环境。"
    )
    add_para(
        doc,
        "本地服务启动后，用户进入页面，系统默认处于 mouse tracking。"
        "这时无需摄像头授权即可加载阅读内容、记录鼠标滚动、更新状态机和显示会话报告。"
        "如果用户点击 gaze 相关入口，系统再检查摄像头权限和安全上下文，进入 WebGazer 初始化与九点校准。"
    )

    add_heading(doc, "12.1.2 阅读任务执行", 3)
    add_para(
        doc,
        "阅读任务开始后，ReadingView 将文本拆分为块并渲染到 ff-reading-content。"
        "PerceptionModule 持续更新行为特征，StateMachine 周期性更新当前状态，DecisionModule 选择策略，"
        "VisualEffects 和 ReadingView 根据策略改变界面。"
        "用户可以在阅读中切换专注模式、查看理解辅助、结束后打开会话报告。"
    )
    add_para(
        doc,
        "LLM 摘要不是阅读任务的前置条件。"
        "如果 server/dev-server.js 检测到 API key，LLMSummaryManager 会启用 LLM 摘要；"
        "如果不可用，系统仍可通过规则摘要、关键词高亮和状态反馈继续完成主要体验。"
        "这种可降级运行方式是课程项目交付的重要可靠性保障。"
    )

    add_heading(doc, "12.2 异常处理", 2)
    add_table(
        doc,
        ["异常场景", "可能原因", "当前代码处理", "验收判断"],
        [
            ["摄像头不可用", "浏览器权限拒绝、设备占用、非安全上下文。", "cameraAccess 检查支持和权限；main.js 可回到 mouse 模式。", "页面不崩溃，仍可完成阅读和状态分析。"],
            ["WebGazer 预测抖动", "光照、头部移动、模型未校准。", "九点校准与 KalmanFilter2D 平滑坐标。", "凝视光标和 block 定位不应剧烈闪动。"],
            ["LLM API 未配置", "缺少 OPENAI_API_KEY 或服务不可达。", "server 返回 disabled/NO_API_KEY；前端 fallback 到规则摘要。", "理解辅助入口仍可使用或明确显示不可用状态。"],
            ["文本导入乱码", "文件编码与浏览器默认解码不一致。", "textEncoding 尝试编码识别和质量评分。", "中文和英文材料应尽量正常显示。"],
            ["状态频繁抖动", "阈值过敏、输入噪声、用户行为边界模糊。", "StateMachine 使用 cooldown、probability threshold 和 hysteresis。", "状态变化应有持续性，不应每秒多次跳变。"],
            ["提示过度打扰", "策略触发太频繁或升级太快。", "InterventionStrategy 设置全局间隔和单策略 cooldown。", "同类提示不会连续密集出现。"],
            ["摘要重复请求", "用户反复进入同一 block。", "LLMSummaryManager 使用 cache 和 prefetch set。", "同段落再次打开时应优先复用缓存。"],
            ["会话结束数据丢失", "页面关闭或用户忘记导出。", "AttentionAnalytics 自动保存历史，SessionReport 支持导出 JSON。", "结束时至少能看到核心统计和时间线。"],
        ],
        [1.15, 1.55, 2.05, 1.9],
        font_size=8.2,
    )

    add_heading(doc, "12.3 验收用例", 2)
    add_para(
        doc,
        "验收不应只看页面是否打开，而应覆盖默认模式、眼动模式、困难辅助、分心恢复、摘要降级和报告导出。"
        "以下用例适合课堂演示前逐项检查，也适合作为答辩时解释系统完整性的依据。"
    )
    add_table(
        doc,
        ["用例", "操作", "预期结果", "涉及模块"],
        [
            ["默认进入阅读", "启动本地服务并打开页面，不授权摄像头。", "系统进入 mouse 模式，阅读内容正常显示，状态可随鼠标和滚动变化。", "index.html；main.js；ReadingView；PerceptionModule"],
            ["导入长文本", "选择一段中文或英文长文本文件。", "文本被正确解码并拆分为阅读块，标题和段落层次基本合理。", "textEncoding；ParagraphSplitter；ReadingView"],
            ["模拟正常阅读", "缓慢滚动并沿文本移动鼠标。", "状态保持或回到 Normal，WPM 和当前 block 正常更新。", "MouseTracker；ScrollAnalyzer；StateMachine；Analytics"],
            ["模拟理解困难", "停在某段超过 8 秒并减少滚动。", "Struggling 概率升高，出现关键词或摘要相关辅助。", "StateMachine；DecisionModule；KeywordExtractor；VisualEffects"],
            ["模拟分心", "长时间不移动鼠标或在 gaze 模式下离开摄像头。", "Distracted 触发轻提醒，持续异常时策略逐步升级。", "FaceDetection；StateMachine；InterventionStrategy"],
            ["模拟恢复", "重新滚动或回到摄像头前继续阅读。", "状态进入 Recovering 或 Normal，显示进度/积极反馈而非继续强提醒。", "StateMachine；VisualEffects；AttentionAnalytics"],
            ["摘要服务可用", "配置 API key 后停留在段落并打开理解辅助。", "摘要通过服务端代理返回，重复打开同段落时使用缓存。", "LLMSummaryManager；dev-server；ReadingView"],
            ["摘要服务不可用", "不配置 API key，执行同样操作。", "系统显示 fallback 或规则摘要，不影响基础阅读流程。", "ParagraphSummarizer；LLMSummaryManager"],
            ["会话报告", "完成阅读后打开 session report。", "显示统计卡片、热力图、时间线，并可导出 JSON。", "AttentionAnalytics；SessionReport"],
            ["语言切换", "切换中英文界面。", "按钮、提示和状态文本随 focusflow-lang-change 更新。", "i18n；index.html；UI 模块"],
        ],
        [1.1, 1.75, 2.05, 1.85],
        font_size=7.9,
    )

    add_heading(doc, "12.4 代码层验收重点", 2)
    add_para(
        doc,
        "代码层验收应重点确认模块之间的接口是否稳定。"
        "PerceptionModule.getFeatures() 的字段需要与 StateMachine 使用的字段一致；"
        "StateMachine 输出的 state.name、confidence、duration 需要被 DecisionModule 正确读取；"
        "InterventionStrategy 输出的 strategy id 需要能映射到 VisualEffects 或 ReadingView 中真实存在的效果；"
        "AttentionAnalytics 记录的状态转换需要与界面状态同步。"
    )
    add_para(
        doc,
        "若要补充自动化测试，可以从纯逻辑模块开始。"
        "StateMachine、AdaptiveThreshold、InterventionStrategy、ParagraphSummarizer 和 KeywordExtractor 都可以在不启动摄像头和浏览器 UI 的情况下进行单元测试。"
        "例如构造 dwellTime > 8000、scrollIdleDuration > 5000 的 features，检查 Struggling 概率是否高于 Normal；"
        "构造连续 Distracted 状态，检查策略是否从 SUBTLE_OVERLAY 逐步升级。"
    )
    add_para(
        doc,
        "UI 层测试则适合用手动验收或浏览器自动化检查。"
        "重点不是验证每个动画像素，而是确认状态变化后对应元素存在、不会遮挡阅读核心内容、冷却时间内不会重复弹出、"
        "摘要卡片能关闭或重新打开、会话报告能在结束后查看。"
        "这些验收点与真实用户体验直接相关。"
    )

    add_heading(doc, "12.5 交付风险控制", 2)
    add_para(
        doc,
        "课程作业交付前，最需要控制的风险是运行环境差异。"
        "不同电脑的摄像头、浏览器权限、Node 版本、网络条件和 API key 配置都可能不同。"
        "FocusFlow 已经通过默认 mouse 模式和摘要 fallback 降低风险，但演示前仍应准备无摄像头路径："
        "直接用鼠标和滚动展示状态变化、困难辅助和会话报告。"
    )
    add_para(
        doc,
        "第二个风险是过度依赖视觉效果。"
        "如果只展示高亮、弹窗和光标，项目容易被看成普通前端动画。"
        "答辩时应把每个效果都对应到状态判断和决策逻辑：高亮来自 Struggling，轻覆盖来自 Distracted，进度反馈来自 Recovering，"
        "会话报告来自 analytics。这样可以证明视觉效果背后有完整的 HCI 逻辑。"
    )
    add_para(
        doc,
        "第三个风险是隐私解释不足。"
        "gaze 模式涉及摄像头，必须强调摄像头是可选增强能力，默认 mouse 模式可以完成核心流程。"
        "摄像头数据用于浏览器端 gaze estimation 和状态辅助，不应被描述为身份识别或长期监控。"
        "这种表述既准确，也能减少用户和评阅者对隐私风险的疑虑。"
    )

    add_heading(doc, "13 结论", 1)
    add_para(
        doc,
        "FocusFlow 实现了一个围绕长文本阅读注意力管理的完整 HCI 原型。"
        "它从浏览器可获得的行为和眼动信号出发，将用户状态解释为 Normal、Distracted、Struggling、Recovering，"
        "再通过自适应策略提供低打扰界面反馈和内容辅助。"
        "项目覆盖了感知、认知、决策、界面反馈、NLP 支持和会话分析，具备课程作业所需的完整系统性。"
    )
    add_para(
        doc,
        "从代码分析看，项目的核心优势是结构层次清楚、状态判断可解释、默认 mouse 模式保证可运行、gaze 模式提供研究扩展、"
        "摘要功能具有 LLM 与规则 fallback、会话报告能够支撑评估。"
        "主要不足是 main.js 和 VisualEffects 的职责较重，状态阈值仍需更多实验数据验证，策略到 UI 的映射还可以进一步显式化。"
    )
    add_para(
        doc,
        "作为人机交互课程作业，FocusFlow 的完成度不仅体现在界面效果，也体现在它对用户状态、交互打扰、隐私权限、个性化阈值和评估指标的综合考虑。"
        "后续若继续推进，可以围绕模块拆分、可访问性、实验评估和策略学习进行迭代，使系统从课程原型进一步接近可长期使用的学习辅助工具。"
    )

    add_heading(doc, "14 参考资料与项目材料", 1)
    add_table(
        doc,
        ["类别", "材料", "用途"],
        [
            ["项目源码", "Vaeloraa/HCI_PROJECT", "课程作业实现、代码分析与演示依据。"],
            ["第三方库", "WebGazer.js", "浏览器端 gaze estimation 与眼动交互基础。"],
            ["第三方库", "Chart.js", "会话数据和图表可视化支持。"],
            ["课程理论", "人机交互中的自适应界面、注意力、可用性评估、隐私与可解释性原则", "用于解释系统设计目标和评估方法。"],
            ["运行环境", "Node.js 本地服务；现代浏览器；可选摄像头；可选 LLM API key", "用于项目运行、演示和摘要增强。"],
        ],
        [1.25, 2.45, 3.05],
        font_size=9.0,
    )


def build_doc():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_section_8(doc)
    add_section_9(doc)
    add_section_10(doc)
    add_section_11(doc)
    add_section_12(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_doc()
    print(path)
