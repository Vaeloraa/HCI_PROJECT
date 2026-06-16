from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "FocusFlow_Project_Report_Expanded_40pages.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEAL = RGBColor(15, 118, 110)
GRAY = RGBColor(86, 96, 110)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
BORDER = "D9DEE7"


def set_run_font(run, size=11, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=size, color=color or RGBColor(0, 0, 0), bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_fixed_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths)))


def add_para(doc, text="", size=11, color=None, bold=False, italic=False, align=None, after=8, before=0, line=1.25):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color or RGBColor(0, 0, 0), bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt({1: 18, 2: 12, 3: 8}.get(level, 8))
    p.paragraph_format.space_after = Pt({1: 10, 2: 6, 3: 4}.get(level, 4))
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), color={1: BLUE, 2: BLUE, 3: DARK_BLUE}.get(level, BLUE), bold=True)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table_width(table, [6.5])
    set_table_borders(table, "D5DCE8")
    set_cell_margins(table, 120, 160, 120, 160)
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=TEAL, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=RGBColor(30, 40, 55))
    add_para(doc, "", after=4)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_fixed_table_width(table, widths)
    set_table_borders(table)
    set_cell_margins(table)
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], HEADER_FILL)
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=NAVY, size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.8 if len(value) > 38 else 9.2)
    add_para(doc, "", after=4)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for idx, size, color in [(1, 16, BLUE), (2, 13, BLUE), (3, 12, DARK_BLUE)]:
        st = styles[f"Heading {idx}"]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("FocusFlow 项目说明书 | Expanded Technical Report")
    set_run_font(r, size=9, color=GRAY, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    r = footer.add_run("FocusFlow - 自适应注意力阅读管理系统")
    set_run_font(r, size=9, color=GRAY)


def cover(doc):
    add_para(doc, "Human-Computer Interaction Project", size=11, color=TEAL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18, before=80)
    add_para(doc, "FocusFlow 项目说明书", size=27, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "自适应注意力阅读管理系统：从多模态感知到自适应阅读干预", size=15, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
    add_para(doc, "扩充版详细文档 / 40+ page edition", size=11, color=GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=60)
    add_callout(doc, "文档定位", "本说明书面向课程答辩、项目归档、后续开发与用户实验准备。它把 FocusFlow 从“可运行原型”解释为一个完整的 HCI 系统：包含用户问题、交互假设、感知管线、认知状态模型、干预策略、界面设计、工程实现、评估方案和伦理边界。")
    add_matrix(doc, ["项目项", "说明"], [
        ["项目名称", "FocusFlow - Adaptive Attention Reading System"],
        ["运行方式", "本地浏览器应用；默认通过 npm start 启动 server/dev-server.js，访问 http://127.0.0.1:8080。"],
        ["核心技术", "WebGazer.js、原生 JavaScript 模块、状态机、鼠标/滚动行为分析、可选 LLM 段落摘要。"],
        ["文档用途", "项目展示、答辩说明、设计与实现归档、实验方案准备、后续维护参考。"],
    ], [1.35, 5.15])
    doc.add_page_break()


sections = [
    ("1. 项目概述", "FocusFlow 是一个面向长文本阅读场景的自适应注意力管理系统。它尝试解决的不是单纯的阅读排版问题，而是阅读过程中更隐蔽的认知状态变化：用户可能仍停留在页面上，却已经从任务中分心；用户可能长时间停在同一段落，不是因为认真阅读，而是因为理解卡住；用户也可能刚刚恢复专注，需要的是低压力的继续提示，而不是强制警告。"),
    ("2. 设计动机", "传统阅读工具主要记录进度、时间和位置，通常不理解这些行为背后的状态含义。FocusFlow 的设计动机是把这些低成本行为信号转化为可解释的状态判断，再用温和、可控、可恢复的交互方式帮助用户回到阅读任务。"),
    ("3. HCI 问题定义", "从人机交互角度看，本项目的核心问题是如何让系统既能感知用户状态，又不把用户变成被监控对象。系统需要在帮助与打扰之间取得平衡：干预太弱会没有效果，干预太强会破坏沉浸感。"),
    ("4. 目标用户与使用场景", "目标用户包括需要阅读长篇资料的学生、研究者、在线课程学习者和需要持续阅读技术文档的开发者。典型场景包括课程论文阅读、技术说明阅读、在线学习材料阅读、答辩前材料复习和需要专注的长文本理解任务。"),
    ("5. 用户痛点分析", "用户在长文本阅读中常见三个痛点：注意力漂移而不自知，遇到难段落时反复停留，以及中断后难以回到原阅读节奏。FocusFlow 把这些痛点拆成可观测行为特征，并进一步映射到状态机。"),
    ("6. 设计原则", "系统遵循五个设计原则：低打扰优先、解释性优先、用户可控优先、降级可用优先、实验可评估优先。这些原则保证系统不是简单地提醒用户“不要分心”，而是提供可解释、可调节、可复盘的阅读辅助。"),
    ("7. 整体系统架构", "项目采用模块化前端结构，主要分为感知层、认知层、决策层、界面层、分析层、NLP 层、国际化层和工具层。每一层承担相对独立职责，便于调试、替换和扩展。"),
    ("8. 感知层概览", "感知层由 PerceptionModule 统一调度，整合 FaceDetection、HeadPoseEstimation、GazeRegionMapping、MouseTracker、ScrollAnalyzer 和 AttentionHeatmap。它的输出不是单个事件，而是可被状态机消费的多模态特征向量。"),
    ("9. 视觉注意力信号", "视觉注意力信号包括人脸是否存在、头部姿态、视线区域、视线置信度和 gaze dwell time。由于浏览器端眼动估计天然存在噪声，系统没有把眼动作为唯一判断依据，而是与鼠标和滚动行为一起融合。"),
    ("10. 鼠标与交互信号", "鼠标信号用于支持无摄像头环境下的演示和使用。系统关注 mouseMoving、idleDuration、clickActivity、horizontalMovementRatio、directionChanges 等信息，用来识别用户是否仍在与阅读内容发生交互。"),
    ("11. 滚动与停留信号", "滚动行为能够反映阅读推进节奏。ScrollAnalyzer 提供 scrollVelocity、scrollIdleDuration、pauseDuration、isScrolling 和 isPaused 等特征，帮助系统区分正常阅读停顿、理解困难和离开任务。"),
    ("12. 段落级注意力热力", "AttentionHeatmap 与段落切分逻辑结合后，可以记录不同段落的停留强度。这为会话报告、理解辅助和后续实验分析提供了基础，也让系统从“页面级”走向“段落级”的解释。"),
    ("13. 特征向量设计", "状态机接收的特征向量包含 facePresent、faceAbsentDuration、interactionActive、dwellTime、scrollVelocity、scrollIdleDuration、trackingMode、faceTracking 等字段。它的意义是把不同来源的数据统一为同一类证据。"),
    ("14. 认知状态模型", "系统定义 Normal、Distracted、Struggling、Recovering 四类状态。Normal 表示稳定阅读；Distracted 表示注意力偏离；Struggling 表示理解停滞；Recovering 表示从中断或困难中恢复。"),
    ("15. 概率式证据融合", "StateMachine 并非只使用单个 if-else 条件，而是为每类状态收集证据并转化为概率式分布。Normal、Distracted、Struggling 和 Recovering 的概率会随输入特征变化，再通过迟滞机制避免频繁跳变。"),
    ("16. 状态转移逻辑", "状态转移同时考虑概率分布、当前状态持续时间和冷却阈值。例如无交互超过阈值可进入 Distracted，长时间停留且不滚动可进入 Struggling，重新交互后可进入 Recovering，稳定一段时间后回到 Normal。"),
    ("17. 阈值与迟滞机制", "系统设置 distractionNoFace、distractionNoInteraction、strugglingDwellTime、recoveringStableTime 和 normalTransitionCooldown 等阈值。迟滞机制让当前状态获得额外稳定性，防止界面因短时噪声闪烁。"),
    ("18. 决策层概览", "Decision 与 InterventionStrategy 的核心任务是把状态判断转化为可执行干预。它不仅看状态名称，还考虑状态置信度、持续时间、连续同状态次数、用户敏感度和策略冷却时间。"),
    ("19. 干预策略库", "策略库包含 Subtle Focus Overlay、Floating Prompt、Sound Alert、Keyword Highlight、Summary Panel、Content Simplification、Progress Indicator 和 Positive Feedback。每种策略都有强度、持续时间、冷却时间和说明。"),
    ("20. 分级干预矩阵", "Distracted 状态可从弱视觉提示升级到浮动提示和声音提醒；Struggling 状态可从关键词高亮升级到摘要面板和内容简化；Recovering 状态更适合进度提示与正向反馈。"),
    ("21. 低打扰策略", "低打扰不是完全不提醒，而是在提醒方式、时机和强度上保持克制。系统通过最小间隔、单策略冷却时间和升级层级，减少同一用户在短时间内被反复打断的风险。"),
    ("22. 阅读界面设计", "界面把复杂推断翻译成用户能理解的状态、进度和提示。核心区域是阅读内容，侧边或仪表区域展示状态、注意力指标、阅读速度、目标进度和事件日志。"),
    ("23. 深度专注模式", "深度专注模式强调减少界面噪声，让用户在需要沉浸阅读时隐藏不必要的信息。这个模式体现了用户控制权：系统可以感知和辅助，但不应强迫用户一直看见分析面板。"),
    ("24. 段落理解辅助", "当用户在段落上长时间停留，系统可以提供关键词、段落概览或摘要面板。这个设计把 Struggling 解释为理解支持需求，而不是简单归因于注意力不足。"),
    ("25. 会话报告", "SessionReport 汇总阅读时长、注意力比例、分心次数、平均恢复时间、段落热力和状态时间线。报告的价值在于复盘，而不是即时惩罚用户。"),
    ("26. 国际化与中文体验", "项目包含 I18n 模块，支持英文与中文界面切换。文档与 PPT 中强调中文体验，是因为课程展示和用户测试需要让按钮、提示、错误建议、报告文本和默认材料保持一致语言。"),
    ("27. 摄像头权限与降级", "WebGazer 依赖摄像头权限，而摄像头在浏览器中会受到协议、权限、设备占用和浏览器兼容性影响。系统提供鼠标追踪作为降级方案，保证项目可演示、可测试、可继续阅读。"),
    ("28. 工程入口与运行方式", "项目入口为 index.html，本地服务由 server/dev-server.js 提供，package.json 中 start 和 serve 都指向本地服务。默认访问地址为 http://127.0.0.1:8080。"),
    ("29. 模块组织", "代码按 perception、cognition、decision、ui、analytics、nlp、i18n 和 utils 分层。这样的组织使 HCI 管线和工程目录保持一致，便于答辩时解释，也便于后续维护。"),
    ("30. 数据流", "典型数据流为：用户行为或摄像头输入进入感知层，感知层输出特征向量，状态机更新状态概率，决策层选择干预策略，界面层呈现反馈，分析层记录事件并生成报告。"),
    ("31. 可调试性", "调试面板、事件日志、状态概率图和段落边界显示帮助开发者理解系统判断来源。对于 HCI 原型而言，可调试性本身也是研究价值，因为它使推断过程更透明。"),
    ("32. NLP 摘要能力", "NLP 层包含 keywordExtractor、paragraphSummarizer 和 llmSummaryManager。摘要功能可以使用本地规则或可选 LLM 接口，目标是在用户理解困难时提供段落级辅助，而不是替代阅读。"),
    ("33. 指标体系", "系统可以记录阅读时间、WPM、注意力比例、分心次数、恢复时间、段落停留、状态分布和干预触发次数。这些指标为后续用户实验提供了可量化依据。"),
    ("34. 用户实验设计", "建议实验设置普通阅读器、鼠标追踪 FocusFlow 和眼动追踪 FocusFlow 三个条件。通过组内或组间设计比较阅读理解、任务时间、恢复时间、主观负担和干预接受度。"),
    ("35. 评估问卷", "主观评估可结合 NASA-TLX、系统可用性量表、感知打扰程度、隐私接受度和系统信任度。问卷应区分“帮助理解”“帮助专注”和“打扰阅读”三类体验。"),
    ("36. 数据记录方案", "实验日志应记录状态时间线、干预类型、触发时刻、段落位置、用户是否关闭提示、会话总结和主观反馈。对于摄像头数据，应避免保存原始视频。"),
    ("37. 伦理与隐私", "注意力系统容易被误解为监控系统，因此必须明确摄像头用途、处理位置、日志保存范围和关闭方式。系统定位应始终是辅助用户恢复注意力，而不是评价用户是否认真。"),
    ("38. 风险与限制", "主要限制包括 WebGazer 精度受设备与光线影响、状态阈值需要实验校准、错误干预可能造成打扰、用户差异会影响判断稳定性，以及 LLM 摘要可能带来延迟和不确定性。"),
    ("39. 后续优化方向", "后续可以加入个性化阈值学习、更多段落级理解辅助、更细粒度的状态解释、实验数据导出、隐私设置面板和更完整的无障碍设计。"),
    ("40. 答辩呈现建议", "答辩时应先展示用户问题，再展示感知到状态再到干预的闭环。PPT 可用封面流程图建立记忆点，用系统架构页解释工程完整性，用评估页说明研究价值。"),
    ("41. 维护手册", "维护时应优先检查本地服务、浏览器权限、摄像头可用性、控制台错误、模块加载顺序、翻译 key 和 WebGazer 初始化。任何涉及状态机的修改都应同时更新阈值说明和实验记录。"),
    ("42. 总结", "FocusFlow 的价值在于把阅读辅助从静态内容呈现推进到动态状态理解。它展示了一个完整 HCI 原型应有的闭环：感知、推断、决策、反馈、记录和评估。"),
]


tables_by_section = {
    7: (["层级", "模块", "职责"], [
        ["感知层", "js/perception", "整合视觉、鼠标、滚动和热力信号，输出多模态特征。"],
        ["认知层", "js/cognition", "维护四类认知状态，计算概率分布并记录状态历史。"],
        ["决策层", "js/decision", "根据状态、严重程度、冷却时间选择干预策略。"],
        ["界面层", "js/ui", "呈现阅读内容、视觉反馈、专注模式、调试面板。"],
        ["分析层", "js/analytics", "生成会话报告、状态时间线和注意力指标。"],
    ], [1.15, 1.65, 3.7]),
    13: (["特征", "来源", "解释"], [
        ["facePresent", "FaceDetection", "表示是否检测到人脸；仅在眼动模式下作为强证据。"],
        ["idleDuration", "MouseTracker", "反映鼠标长期无动作，可支持分心判断。"],
        ["scrollIdleDuration", "ScrollAnalyzer", "反映阅读推进停止时间，用于识别暂停和困难。"],
        ["dwellTime", "GazeRegion / Scroll", "表示在当前区域停留时长，是理解停滞的重要信号。"],
    ], [1.6, 1.7, 3.2]),
    20: (["状态", "轻度", "中度", "重度"], [
        ["Distracted", "视觉弱提示", "浮动提示", "声音提醒 + 强视觉线索"],
        ["Struggling", "关键词高亮", "段落摘要", "内容简化 / 解释"],
        ["Recovering", "进度提示", "正向反馈", "降低干预频率"],
        ["Normal", "不干预", "记录指标", "保持沉浸"],
    ], [1.25, 1.75, 1.75, 1.75]),
    34: (["实验条件", "输入方式", "比较目的"], [
        ["普通阅读器", "无状态感知", "建立基础阅读表现基线。"],
        ["鼠标追踪 FocusFlow", "鼠标、滚动、停留", "验证低成本信号是否足以改善恢复。"],
        ["眼动追踪 FocusFlow", "摄像头、眼动、鼠标、滚动", "验证多模态信号是否提升状态判断。"],
    ], [1.7, 2.1, 2.7]),
    38: (["风险", "影响", "缓解方式"], [
        ["摄像头不可用", "眼动模式无法启动", "降级到鼠标追踪并给出可操作错误提示。"],
        ["误判分心", "打断阅读流", "提高冷却时间、使用弱提示、允许关闭干预。"],
        ["摘要不准确", "影响理解", "标注摘要为辅助信息，保留原文为主。"],
        ["隐私担忧", "降低信任", "透明说明用途，不保存原始视频，提供关闭入口。"],
    ], [1.45, 1.85, 3.2]),
}


def add_section_page(doc, index, title, body):
    add_heading(doc, title, 1)
    add_para(doc, body, size=11, color=RGBColor(20, 30, 40), line=1.333, after=8)
    add_para(doc, "详细说明", size=11.2, color=TEAL, bold=True, after=4)
    add_para(
        doc,
        f"在本项目中，{title.split(' ', 1)[-1]}不是孤立功能，而是与阅读任务、状态推断、干预策略和会话复盘相互连接。设计时需要同时考虑用户是否愿意接受、开发者是否能够调试、实验者是否能够测量，以及系统在设备条件不稳定时是否仍然可用。",
        size=10.8,
        color=RGBColor(25, 35, 45),
        line=1.3,
        after=7,
    )
    add_para(
        doc,
        "实现关注点包括：输入信号的稳定性、状态解释的清晰度、干预触发的节制性、界面反馈的可理解性、日志记录的可复盘性，以及用户隐私与控制权。若后续继续开发，应优先补齐实验数据导出、阈值个性化、错误提示可操作性和段落级辅助质量评估。",
        size=10.8,
        color=RGBColor(25, 35, 45),
        line=1.3,
        after=7,
    )
    add_callout(doc, "章节要点", f"{title} 这一部分的核心价值，是把项目从功能描述推进到可解释的 HCI 设计论证。读者应能看到该模块为什么存在、依赖哪些信号、如何影响用户体验，以及如何在实验中被验证。")
    if index in tables_by_section:
        headers, rows, widths = tables_by_section[index]
        add_matrix(doc, headers, rows, widths)
    add_para(doc, "检查问题", size=10.8, color=DARK_BLUE, bold=True, after=4)
    for q in [
        "该部分是否能被用户理解，而不是只被开发者理解？",
        "该部分是否提供了可观测、可记录、可复盘的依据？",
        "如果输入信号不稳定，系统是否仍有可用的降级路径？",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(q)
        set_run_font(r, size=10.3, color=RGBColor(35, 45, 60))
    doc.add_page_break()


def appendix(doc):
    add_heading(doc, "附录 A：运行与检查清单", 1)
    add_matrix(doc, ["步骤", "检查内容", "通过标准"], [
        ["1", "运行 npm start", "本地服务启动且端口 8080 可访问。"],
        ["2", "打开 http://127.0.0.1:8080", "页面加载，无关键资源 404。"],
        ["3", "切换鼠标追踪", "无需摄像头即可看到状态变化。"],
        ["4", "切换眼动追踪", "浏览器请求摄像头权限，预览与校准流程可用。"],
        ["5", "导入文本", "段落被识别，阅读区域可滚动。"],
        ["6", "查看会话报告", "包含时间、注意力比例、状态时间线和段落热力。"],
    ], [0.7, 2.7, 3.1])
    add_heading(doc, "附录 B：答辩讲述顺序", 1)
    for item in [
        "先说明长文本阅读中的注意力漂移和理解停滞问题。",
        "再展示 FocusFlow 的闭环：感知、状态、决策、反馈、复盘。",
        "随后解释系统架构与核心模块，突出模块化和可扩展性。",
        "演示鼠标模式，保证现场不依赖摄像头权限。",
        "如果环境允许，再演示眼动模式和摄像头校准。",
        "最后用评估方案和伦理边界说明项目的研究完整性。",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=RGBColor(25, 35, 45))


def main():
    doc = Document()
    configure_doc(doc)
    cover(doc)
    add_heading(doc, "目录式章节概览", 1)
    for title, _ in sections:
        add_para(doc, title, size=10.3, color=DARK_BLUE, bold=True, after=2)
    doc.add_page_break()

    for idx, (title, body) in enumerate(sections, start=1):
        add_section_page(doc, idx, title, body)

    appendix(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
